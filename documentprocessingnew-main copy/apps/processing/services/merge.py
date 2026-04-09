from copy import deepcopy
import json
import logging
from io import BytesIO

import fitz
from bs4 import BeautifulSoup

from django.db import transaction
from django.utils import timezone
from django.core.files.base import ContentFile

from apps.documents.models import Document
from apps.processing.models import SubmittedPage, MergedDocument, ApprovedDocument
from common.enums import PipelineStatus, ReviewStatus, MergeStatus, DocumentStatus

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

try:
    from docxcompose.composer import Composer
except Exception:  # pragma: no cover
    Composer = None


class MergeService:
    @staticmethod
    def _load_submission_pdf_bytes(submission: SubmittedPage):
        """
        Return raw PDF bytes from the submitted artifact when available.
        Non-PDF artifacts are ignored here (fallback to re-bake).
        """
        if not submission.output_page_file:
            return None

        try:
            name = (submission.output_page_file.name or "").lower()
            submission.output_page_file.seek(0)
            content = submission.output_page_file.read()
        except Exception:
            return None

        if not content:
            return None

        # Accept explicit .pdf files or any file that actually contains PDF bytes.
        if name.endswith(".pdf") or content[:4] == b"%PDF":
            return content
        return None

    @staticmethod
    def _load_submission_docx_bytes(submission: SubmittedPage):
        """
        Return raw DOCX bytes from the submitted artifact when available.
        Accept .docx files or byte streams that look like DOCX zip containers.
        """
        if not submission.output_page_file:
            return None

        try:
            name = (submission.output_page_file.name or "").lower()
            submission.output_page_file.seek(0)
            content = submission.output_page_file.read()
        except Exception:
            return None

        if not content:
            return None

        # DOCX files are ZIP containers and usually start with PK.
        if name.endswith(".docx") or content[:2] == b"PK":
            return content
        return None

    @staticmethod
    def _get_original_pdf_page_as_docx(document, page_number):
        """
        Converts a single page from the original source PDF into a minimal DOCX
        containing an embedded image of that page. Used as a fallback for pages
        that have not been assigned, edited, or submitted.
        """
        import io
        try:
            import fitz
            from docx.shared import Inches
            from pathlib import Path as _Path

            # Resolve original PDF path
            src = None
            if getattr(document, 'file', None):
                try:
                    src = document.file.path
                except Exception:
                    pass
            if not src and hasattr(document, 'original_file') and document.original_file:
                try:
                    src = document.original_file.path
                except Exception:
                    pass

            if not src or not _Path(src).exists():
                raise ValueError(f"Original PDF not found for document {document.id}")

            # Open PDF with PyMuPDF
            pdf_doc = fitz.open(src)
            if page_number < 1 or page_number > len(pdf_doc):
                raise ValueError(f"Page {page_number} out of range (PDF has {len(pdf_doc)} pages)")

            # Render page to a high quality PNG
            page = pdf_doc[page_number - 1]
            zoom = 2.0  # Higher resolution for the embedded image
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img_stream = io.BytesIO(img_bytes)
            pdf_doc.close()

            # Create a DOCX and embed the image
            doc = DocxDocument()
            
            # Adjust page margins to be small
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            import docx

            # Add a small note
            para = doc.add_paragraph()
            run = para.add_run(f"[Original Page {page_number} — Unedited]")
            run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)

            # Add the image
            doc.add_picture(img_stream, width=Inches(7.0))

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

        except Exception as exc:
            logger.warning(f"Could not extract original PDF page {page_number} for document {document.id}: {exc}")
            # Last resort: blank placeholder docx
            doc = DocxDocument()
            doc.add_paragraph(f"[Page {page_number} — content unavailable]")
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

    @staticmethod
    def _effective_page_number(submission: SubmittedPage):
        """
        Resolve the page number from the submission itself or from the related page object.
        """
        page_number = getattr(submission, "page_number", None)
        if page_number is None and getattr(submission, "page", None):
            page_number = getattr(submission.page, "page_number", None)
        return page_number

    @staticmethod
    def _submission_sort_key(submission: SubmittedPage):
        """
        Sort by page number ascending, then by submitted_at descending, then by id descending.
        This lets the latest submission win for duplicate page numbers.
        """
        page_number = MergeService._effective_page_number(submission)
        if page_number is None:
            page_number = 10**9

        submitted_at = getattr(submission, "submitted_at", None)
        submitted_ts = submitted_at.timestamp() if submitted_at else 0

        submission_id = getattr(submission, "id", 0) or 0
        return (page_number, -submitted_ts, -submission_id)

    @staticmethod
    def _collect_best_versions_ledger(document: Document):
        """
        Deterministic Resolver (ZeroLossMergeEngine):
        1. Select ONLY records where review_status == APPROVED.
        2. Group by page_number.
        3. Within each group, pick the LATEST version (ORDER BY submitted_at DESC, id DESC).
        """
        approved_qs = (
            SubmittedPage.objects.filter(
                document=document,
                review_status=ReviewStatus.APPROVED,
            )
            .select_related("submitted_by", "page", "assignment")
            .all()
        )

        submissions = list(approved_qs)
        # Sort key: Page Number (ASC), submitted_at (DESC), id (DESC)
        submissions.sort(key=lambda s: (
            MergeService._effective_page_number(s) or 0,
            -(s.submitted_at.timestamp() if s.submitted_at else 0),
            -s.id
        ))

        page_map = {}
        for submission in submissions:
            page_number = MergeService._effective_page_number(submission)
            if page_number not in page_map:
                page_map[page_number] = submission

        return page_map

    @staticmethod
    def _validate_submission_record(submission: SubmittedPage, page_number: int):
        """
        Verify that the stored submission has the required traceability fields.
        """
        missing = []

        assignment_id = getattr(submission, "assignment_id", None)
        if assignment_id is None and getattr(submission, "assignment", None) is not None:
            assignment_id = getattr(submission.assignment, "id", None)
        if assignment_id is None:
            missing.append("assignment_id")

        page_id = getattr(submission, "page_id", None)
        if page_id is None and getattr(submission, "page", None) is not None:
            page_id = getattr(submission.page, "id", None)
        if page_id is None:
            missing.append("page_id")

        submitted_by_id = getattr(submission, "submitted_by_id", None)
        if submitted_by_id is None and getattr(submission, "submitted_by", None) is not None:
            submitted_by_id = getattr(submission.submitted_by, "id", None)
        if submitted_by_id is None:
            missing.append("user_id")

        if getattr(submission, "submitted_at", None) is None:
            missing.append("timestamp")

        if missing:
            raise ValueError(
                f"Integrity Error: Page {page_number} is missing: {', '.join(missing)}"
            )

    @staticmethod
    def _append_docx_body(target_doc, source_doc):
        """
        Fallback DOCX merge helper when docxcompose is unavailable.
        Appends body XML nodes in order. This preserves paragraphs and tables.
        """
        target_body = target_doc._element.body
        source_body = source_doc._element.body

        for child in source_body:
            # Skip the source document section properties.
            if child.tag.endswith("sectPr"):
                continue
            target_body.append(deepcopy(child))

    @staticmethod
    def _merge_docx_bytes(docx_blobs):
        """
        Merge multiple DOCX byte blobs into one DOCX buffer.
        Uses docxcompose if installed; otherwise falls back to XML body append.
        """
        if not docx_blobs:
            raise ValueError("No DOCX content found to merge.")

        if DocxDocument is None:
            raise RuntimeError(
                "python-docx is not installed or failed to import. DOCX merge cannot run."
            )

        if Composer is not None:
            base_doc = DocxDocument(BytesIO(docx_blobs[0]))
            composer = Composer(base_doc)

            for blob in docx_blobs[1:]:
                composer.append(DocxDocument(BytesIO(blob)))

            buffer = BytesIO()
            composer.save(buffer)
            buffer.seek(0)
            return buffer

        merged_doc = DocxDocument(BytesIO(docx_blobs[0]))
        for blob in docx_blobs[1:]:
            source_doc = DocxDocument(BytesIO(blob))
            MergeService._append_docx_body(merged_doc, source_doc)

        buffer = BytesIO()
        merged_doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def _generate_single_page_docx_bytes(document: Document, page_obj):
        """
        Fallback generator for a single page when the saved edited DOCX artifact
        is missing. This reuses the existing export pipeline.
        """
        from apps.processing.services.export import ExportService

        page_export = ExportService.generate_word_export(document, pages=[page_obj])

        if hasattr(page_export, "getvalue"):
            data = page_export.getvalue()
            if data:
                return data

        if hasattr(page_export, "read"):
            data = page_export.read()
            if data:
                return data

        if isinstance(page_export, bytes):
            return page_export

        raise ValueError("ExportService did not return DOCX bytes for a page export.")

    @staticmethod
    def merge_approved_pages(document: Document, admin_user_id=None):
        """
        Takes all APPROVED SubmittedPage records for a Document and merges them
        into a final PDF. (Section 10).
        Creates MergedDocument and ApprovedDocument records.
        """
        with transaction.atomic():
            document = Document.objects.select_for_update().get(id=document.id)

            # 1. Validation check
            if document.pipeline_status == PipelineStatus.APPROVED and document.final_file:
                return

            total_pages = document.total_pages or 0
            if total_pages == 0:
                raise ValueError("Document has no pages recorded.")

            approved_qs = SubmittedPage.objects.filter(
                document=document,
                review_status=ReviewStatus.APPROVED
            ).order_by('page_number', '-submitted_at', '-id').distinct('page_number')

            approved_count = approved_qs.count()
            if approved_count != total_pages:
                found_pages = set(approved_qs.values_list('page_number', flat=True))
                expected_pages = set(range(1, total_pages + 1))
                missing = sorted(list(expected_pages - found_pages))
                raise ValueError(
                    f"Cannot merge: {approved_count}/{total_pages} pages approved. Missing: {missing}"
                )

            merged_doc, _ = MergedDocument.objects.get_or_create(document=document)

            try:
                doc_pdf = fitz.open()

                for page_num in range(1, total_pages + 1):
                    submission = approved_qs.filter(page_number=page_num).first()

                    if not submission:
                        raise ValueError(
                            f"Integrity Error: Approved submission for page {page_num} went missing during merge."
                        )

                    page_pdf_bytes = MergeService._load_submission_pdf_bytes(submission)
                    if page_pdf_bytes is None:
                        from apps.processing.services.pdf_baking import PDFBakeService
                        try:
                            page_pdf_bytes = PDFBakeService.bake_page_edits(submission.page)
                        except Exception as bake_err:
                            logger.error(
                                f"Re-bake failed for page {page_num} during merge: {bake_err}"
                            )
                            raise ValueError(
                                f"Page {page_num} is approved but has no mergeable PDF artifact."
                            )

                    with fitz.open(stream=page_pdf_bytes, filetype="pdf") as page_pdf:
                        doc_pdf.insert_pdf(page_pdf)

                result_buffer = BytesIO()
                doc_pdf.save(result_buffer)
                doc_pdf.close()

                filename = f"final_merged_{document.doc_ref}.pdf"
                merged_doc.merged_file.save(filename, ContentFile(result_buffer.getvalue()), save=False)

                merged_doc.merge_status = MergeStatus.COMPLETED
                merged_doc.merge_completed_at = timezone.now()
                merged_doc.merged_by_id = admin_user_id
                merged_doc.save()

                ApprovedDocument.objects.update_or_create(
                    document=document,
                    defaults={
                        'merged_document': merged_doc,
                        'approved_by_id': admin_user_id,
                        'approval_notes': "Auto-generated upon completion of all page reviews."
                    }
                )

                from apps.audit.models import AuditLog
                from common.enums import AuditEventType
                AuditLog.objects.create(
                    action=AuditEventType.DOC_COMPLETED,
                    document_id=document.id,
                    actor_id=admin_user_id,
                    metadata={'page_count': total_pages}
                )

                document.final_file.save(filename, ContentFile(result_buffer.getvalue()), save=False)
                document.pipeline_status = PipelineStatus.MERGED
                document.status = DocumentStatus.COMPLETED
                document.completed_at = timezone.now()
                document.save(update_fields=['final_file', 'pipeline_status', 'status', 'completed_at'])

            except Exception as e:
                document.pipeline_status = PipelineStatus.FAILED
                document.pipeline_error = f"Merge failed: {str(e)}"
                document.save(update_fields=['pipeline_status', 'pipeline_error'])
                logger.error(f"Error merging document {document.id}: {e}")
                raise e

    @staticmethod
    def merge_approved_docx_pages(document: Document, admin_user_id=None):
        """
        Merge engine that concatenates the full DOCX payloads of each consecutive 
        AssignmentBundle in strictly ascending order. No per-page mapping or versioning.
        """
        import hashlib
        from io import BytesIO
        from django.core.files.base import ContentFile
        from docx import Document as DocxDocument
        from apps.desktop_bridge.models import AssignmentBundle, UploadedPDF

        try:
            from docxcompose.composer import Composer
        except ImportError:
            Composer = None

        with transaction.atomic():
            document = Document.objects.select_for_update().get(id=document.id)
            merged_rec, _ = MergedDocument.objects.get_or_create(document=document)

            if document.final_word_file:
                try:
                    document.final_word_file.open('rb')
                    document.final_word_file.close()
                    logger.info(f"BundleMergeEngine: Document {document.id} already merged. Returning existing file.")
                    return document.final_word_file.url
                except Exception:
                    pass

            merged_rec.merge_attempt_count += 1
            merged_rec.merge_status = MergeStatus.PENDING
            merged_rec.merge_started_at = timezone.now()
            merged_rec.save()

            try:
                # 1) Get all AssignmentBundles strictly in page_start order (correct page sequence)
                bundles = AssignmentBundle.objects.filter(document=document).order_by('page_start')
                if not bundles.exists():
                    raise ValueError(f"BundleMergeEngine Error: Document {document.id} has no bundles.")

                ordered_docx_blobs = []
                manifest_pages = []
                
                # 2) For each bundle, collect its DOCX content — try multiple sources
                for bundle in bundles:

                    # ── STAGE 1a: Atomic output_file (new strict flow) ──
                    blob = None
                    source_label = None

                    for file_field, label in [
                        (bundle.output_file, 'output_file'),
                        (bundle.result_pdf,  'result_pdf'),
                    ]:
                        if not file_field:
                            continue
                        try:
                            file_field.open('rb')
                            candidate = file_field.read()
                            file_field.close()
                            if candidate and len(candidate) > 64:  # skip empty/stub files
                                blob = candidate
                                source_label = label
                                break
                        except Exception as ex:
                            logger.warning(f"BundleMergeEngine: Cannot read {label} for bundle {bundle.id}: {ex}")

                    if blob:
                        ordered_docx_blobs.append(blob)
                        manifest_pages.append({
                            "bundle_id": str(bundle.id),
                            "type": f"atomic_bundle:{source_label}",
                            "page_start": bundle.page_start,
                            "page_end": bundle.page_end
                        })
                        logger.info(f"BundleMergeEngine: Bundle {bundle.bundle_index} using {source_label} ({bundle.page_start}-{bundle.page_end})")
                        continue

                    # ── STAGE 1b: Check UploadedPDF records for this bundle ──
                    latest_upload = UploadedPDF.objects.filter(bundle=bundle).order_by('-created_at').first()
                    if latest_upload and latest_upload.file:
                        try:
                            latest_upload.file.open('rb')
                            candidate = latest_upload.file.read()
                            latest_upload.file.close()
                            if candidate and len(candidate) > 64:
                                ordered_docx_blobs.append(candidate)
                                manifest_pages.append({
                                    "bundle_id": str(bundle.id),
                                    "type": "uploaded_pdf",
                                    "page_start": bundle.page_start,
                                    "page_end": bundle.page_end
                                })
                                logger.info(f"BundleMergeEngine: Bundle {bundle.bundle_index} using UploadedPDF {latest_upload.id}")
                                continue
                        except Exception as ex:
                            logger.warning(f"BundleMergeEngine: Cannot read UploadedPDF for bundle {bundle.id}: {ex}")

                    # ── STAGE 2: Fallback to per-page SubmittedPage records ──
                    from apps.processing.models import SubmittedPage, ReviewStatus
                    bundle_pages = range(bundle.page_start, bundle.page_end + 1)
                    page_blobs = []
                    for page_num in bundle_pages:
                        submission = SubmittedPage.objects.filter(
                            bundle=bundle,
                            page_number=page_num,
                            review_status=ReviewStatus.APPROVED
                        ).order_by('-submitted_at', '-id').first()

                        if not submission or not submission.output_page_file:
                            raise ValueError(
                                f"BundleMergeEngine Error: Bundle {bundle.bundle_index} "
                                f"(pages {bundle.page_start}-{bundle.page_end}) has no uploaded file "
                                f"and no approved submission for page {page_num}. "
                                f"Please upload the bundle DOCX file via the desktop agent."
                            )

                        submission.output_page_file.open('rb')
                        page_blob = submission.output_page_file.read()
                        submission.output_page_file.close()
                        if not page_blob:
                            raise ValueError(f"BundleMergeEngine Error: empty file for page {page_num} in bundle {bundle.id}")
                        page_blobs.append(page_blob)
                        manifest_pages.append({
                            "bundle_id": str(bundle.id),
                            "page_number": page_num,
                            "submission_id": str(submission.id)
                        })

                    ordered_docx_blobs.extend(page_blobs)


                # 3) Merge the DOCX blobs
                logger.info(f"BundleMergeEngine: Merging {len(ordered_docx_blobs)} bundle payloads.")
                if Composer and len(ordered_docx_blobs) > 1:
                    master = DocxDocument(BytesIO(ordered_docx_blobs[0]))
                    composer = Composer(master)
                    for i in range(1, len(ordered_docx_blobs)):
                        composer.append(DocxDocument(BytesIO(ordered_docx_blobs[i])))
                    
                    final_docx_buffer = BytesIO()
                    master.save(final_docx_buffer)
                else:
                    final_docx_buffer = MergeService._merge_docx_bytes(ordered_docx_blobs)
                
                final_docx_buffer.seek(0)
                
                # 4) Output & Store
                timestamp_str = timezone.now().strftime("%Y%m%d_%H%M%S")
                final_filename = f"Final_Merged_{document.id}_{timestamp_str}.docx"
                
                merged_rec.merged_file.save(final_filename, ContentFile(final_docx_buffer.getvalue()), save=False)
                merged_rec.merge_status = MergeStatus.COMPLETED
                merged_rec.merge_completed_at = timezone.now()
                merged_rec.manifest_json = {
                    "document_id": str(document.id),
                    "merged_at": merged_rec.merge_completed_at.isoformat(),
                    "total_bundles": len(bundles),
                    "bundles_merged": manifest_pages
                }
                merged_rec.save()

                document.final_word_file.save(final_filename, ContentFile(final_docx_buffer.getvalue()), save=False)
                document.final_word_generated_at = timezone.now()
                document.final_word_manifest = merged_rec.manifest_json
                document.pipeline_status = PipelineStatus.MERGED
                document.status = DocumentStatus.COMPLETED
                document.completed_at = timezone.now()
                document.save()

                logger.info(f"✔ BundleMergeEngine Success: {final_filename}")
                return document.final_word_file.url

            except Exception as e:
                error_summary = f"BundleMergeEngine Failure: {str(e)}"
                logger.error(f"❌ {error_summary}", exc_info=True)
                
                document.pipeline_status = PipelineStatus.FAILED
                document.pipeline_error = error_summary
                document.save(update_fields=['pipeline_status', 'pipeline_error'])
                
                merged_rec.merge_status = MergeStatus.FAILED
                merged_rec.merge_error = error_summary
                merged_rec.save(update_fields=['merge_status', 'merge_error'])
                raise