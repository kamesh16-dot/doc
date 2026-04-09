from __future__ import annotations

import logging
import re
from collections import defaultdict
from dataclasses import dataclass, field
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from apps.documents.models import Block, Page, PageImage, PageTable
from apps.processing.services.corrector import TextCorrector
from common.enums import ReviewStatus

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Layout constants
# ---------------------------------------------------------------------------
BASE_MARGIN_PT = 36.0                 # 0.5 inch
HEADER_DIST_PT = 18.0
FOOTER_DIST_PT = 18.0

DEFAULT_PAGE_W_PT = 595.0             # A4 width in PDF points
DEFAULT_PAGE_H_PT = 842.0             # A4 height in PDF points

# RELAXED TOLERANCES FOR FINANCIAL STATEMENTS
Y_SNAP_TOLERANCE_PT = 5.0             # Increased to group TOC and Balance Sheets properly
ROW_GROUP_TOLERANCE_PT = 5.0          # Increased for better row clustering
CENTER_TOLERANCE_PT = 50.0            # Greatly increased to catch slightly off-center headers
RIGHT_TOLERANCE_PT = 50.0             # Increased to catch right-aligned page numbers/$ amounts
X_POSITION_TOLERANCE_PT = 4.0         # Tolerance for overlap
MARGIN_SNAP_TOLERANCE = 12.0          # If x is within 12pt of margin, snap it to 0 indent

LINE_SPACING_DEFAULT = 1.0            # Set to 1.0 for tighter financial statement look
LINE_SPACING_TIGHT = 0.10
MAX_SPACER_GAP_PT = 800.0             # Max vertical gap 
MIN_COLUMN_WIDTH_PT = 10.0

FONT_MAP = {
    "arial": "Arial", "helv": "Arial", "helvetica": "Arial",
    "times": "Times New Roman", "serif": "Times New Roman",
    "courier": "Courier New", "mono": "Courier New", "consolas": "Courier New",
    "georgia": "Georgia", "verdana": "Verdana", "tahoma": "Tahoma",
    "trebuchet": "Trebuchet MS", "comic": "Comic Sans MS",
    "impact": "Impact", "symbol": "Symbol",
}

# ---------------------------------------------------------------------------
# Data containers
# ---------------------------------------------------------------------------
@dataclass
class LayoutItem:
    kind: str                      
    obj: Any
    x: float
    y: float
    w: float
    h: float
    z: int = 0
    order: int = 0
    original_x: float = 0.0        
    alignment: Optional[str] = None  

    @property
    def bottom(self) -> float: return self.y + self.h

    @property
    def right(self) -> float: return self.x + self.w

    @property
    def center_x(self) -> float: return self.x + self.w / 2.0

@dataclass
class LayoutRow:
    items: List[LayoutItem] = field(default_factory=list)
    y: float = 0.0
    
    @property
    def top(self) -> float: return min(item.y for item in self.items) if self.items else 0.0
    
    @property
    def bottom(self) -> float: return max(item.bottom for item in self.items) if self.items else 0.0

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _safe_float(value: Any, default: float = 0.0) -> float:
    try: return float(value) if value is not None else default
    except Exception: return default

def _safe_int(value: Any, default: int = 0) -> int:
    try: return int(value) if value is not None else default
    except Exception: return default

def _parse_color(hex_str: str) -> Optional[RGBColor]:
    if not hex_str: return None
    s = str(hex_str).strip()
    m = re.match(r"rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", s, re.I)
    if m:
        try: return RGBColor(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except Exception: return None

    named = {
        "black": "000000", "white": "FFFFFF", "red": "FF0000",
        "green": "008000", "blue": "0000FF", "gray": "808080", "grey": "808080",
    }
    s = named.get(s.lower(), s).lstrip("#")
    if len(s) == 3: s = "".join(c * 2 for c in s)
    if len(s) != 6: return None
    try: return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except Exception: return None

def _map_font(raw_name: str) -> str:
    lower = (raw_name or "").lower()
    for key, mapped in FONT_MAP.items():
        if key in lower: return mapped
    return "Calibri"

def _pt_to_twips(pt: float) -> int:
    return int(pt * 20)

def _set_section_geometry(section, page_w_pt: float, page_h_pt: float) -> None:
    section.page_width = Pt(page_w_pt)
    section.page_height = Pt(page_h_pt)
    section.top_margin = Pt(BASE_MARGIN_PT)
    section.bottom_margin = Pt(BASE_MARGIN_PT)
    section.left_margin = Pt(BASE_MARGIN_PT)
    section.right_margin = Pt(BASE_MARGIN_PT)
    section.header_distance = Pt(HEADER_DIST_PT)
    section.footer_distance = Pt(FOOTER_DIST_PT)

def _zero_table_cell_margins(table) -> None:
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._element.insert(0, tbl_pr)
    
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is not None: tbl_pr.remove(tbl_cell_mar)
    
    tbl_cell_mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        tbl_cell_mar.append(m)
    tbl_pr.append(tbl_cell_mar)

def _remove_all_borders(table) -> None:
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._element.insert(0, tbl_pr)
    
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is not None: tbl_pr.remove(tbl_borders)
    
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)
    
    if table.rows:
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = tc_pr.find(qn("w:tcBorders"))
                if tc_borders is not None: tc_pr.remove(tc_borders)
                tc_borders = OxmlElement("w:tcBorders")
                for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    b = OxmlElement(f"w:{s}")
                    b.set(qn("w:val"), "none")
                    b.set(qn("w:sz"), "0")
                    b.set(qn("w:space"), "0")
                    b.set(qn("w:color"), "auto")
                    tc_borders.append(b)
                tc_pr.append(tc_borders)

def _apply_borderless_style(table) -> None:
    table.style = 'Normal Table'
    _zero_table_cell_margins(table)
    _remove_all_borders(table)

def _apply_underline(run, block: Block) -> None:
    run.underline = bool(getattr(block, "is_underlined", False))

def _set_table_indent(table, x_pt: float) -> None:
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._element.insert(0, tbl_pr)
    
    existing = tbl_pr.find(qn("w:tblInd"))
    if existing is not None: tbl_pr.remove(existing)
    
    indent_pt = max(0.0, x_pt - BASE_MARGIN_PT)
    if indent_pt < MARGIN_SNAP_TOLERANCE: 
        indent_pt = 0.0 # Snap to margin
        
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(_pt_to_twips(indent_pt)))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

def _set_table_width(table, width_pt: float) -> None:
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._element.insert(0, tbl_pr)
    
    existing = tbl_pr.find(qn("w:tblW"))
    if existing is not None: tbl_pr.remove(existing)
    
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(_pt_to_twips(width_pt)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

def _set_column_width(column, width_pt: float) -> None:
    column.width = Pt(max(MIN_COLUMN_WIDTH_PT, width_pt))

def _set_cell_width(cell, width_pt: float) -> None:
    width_pt = max(MIN_COLUMN_WIDTH_PT, width_pt)
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcW"))
    if existing is not None: tc_pr.remove(existing)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(_pt_to_twips(width_pt)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)

def _add_spacer_paragraph(doc, space_before: float) -> None:
    if space_before <= 0: return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(min(space_before, MAX_SPACER_GAP_PT))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

def _detect_list_kind(text: str) -> Optional[str]:
    if not text: return None
    t = text.strip()
    if re.match(r"^(\u2022|\-|\*|–|—)\s+", t): return "bullet"
    if re.match(r"^\(?\d+[\.\)]\s+", t): return "number"
    if re.match(r"^\(?[a-zA-Z][\.\)]\s+", t): return "alpha"
    return None

def _clean_text(text: str) -> str:
    if not text: return ""
    text = TextCorrector.process_block(text)
    text = re.sub(r"[ \t]+", " ", text)
    for ph in ("[Line Area]", "[Image Area]", "[Table Cell]", "[New Table]"):
        if text.strip() == ph: return ""
    return re.sub(r"^[>•\-*]\s*", "", text.strip()).strip()

def _detect_alignment_from_position(x: float, w: float, page_width_pt: float, content_width_pt: float) -> str:
    if w <= 0: return "left"
    
    item_center = x + (w / 2.0)
    item_end = x + w
    page_center = page_width_pt / 2.0
    content_end = page_width_pt - BASE_MARGIN_PT
    
    if abs(item_center - page_center) < CENTER_TOLERANCE_PT: return "center"
    if abs(item_end - content_end) < RIGHT_TOLERANCE_PT: return "right"
    return "left"

# ---------------------------------------------------------------------------
# Export Service
# ---------------------------------------------------------------------------
class ExportService:

    @staticmethod
    def generate_word_export(document, include_unapproved: bool = True, pages=None) -> BytesIO:
        doc = DocxDocument()
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        if not doc.sections: doc.add_section()
        _set_section_geometry(doc.sections[0], DEFAULT_PAGE_W_PT, DEFAULT_PAGE_H_PT)

        export_pages = ExportService._resolve_pages(document, include_unapproved, pages)

        for page_index, page in enumerate(export_pages):
            page_w = _safe_float(getattr(page, "pdf_page_width", None), DEFAULT_PAGE_W_PT)
            page_h = _safe_float(getattr(page, "pdf_page_height", None), DEFAULT_PAGE_H_PT)

            if page_index > 0:
                _set_section_geometry(doc.add_section(), page_w, page_h)

            ExportService._render_page(doc, page, page_w, page_h)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def _resolve_pages(document, include_unapproved: bool, pages: Optional[List]) -> List[Page]:
        if pages is not None: return list(pages)
        qs = document.pages.all().order_by("page_number")
        
        try: pages_list = list(qs) if include_unapproved else list(qs.filter(status=ReviewStatus.APPROVED.value))
        except Exception: pages_list = list(qs.filter(status="APPROVED"))

        if pages_list:
            dw = _safe_float(getattr(pages_list[0], "pdf_page_width", None), DEFAULT_PAGE_W_PT)
            dh = _safe_float(getattr(pages_list[0], "pdf_page_height", None), DEFAULT_PAGE_H_PT)
            for p in pages_list:
                if getattr(p, "pdf_page_width", None) is None: setattr(p, "pdf_page_width", dw)
                if getattr(p, "pdf_page_height", None) is None: setattr(p, "pdf_page_height", dh)

        return pages_list

    @staticmethod
    def _render_page(doc, page: Page, page_width_pt: float, page_height_pt: float) -> None:
        items = ExportService._collect_layout_items(page, page_width_pt)
        if not items: return

        content_width_pt = page_width_pt - (2 * BASE_MARGIN_PT)
        table_groups, non_table_items = ExportService._split_tables(items)
        render_queue: List[Dict[str, Any]] = []

        for table_id, blocks in table_groups.items():
            render_queue.append({
                "type": "table", 
                "y": min(_safe_float(getattr(b, "y", 0.0)) for b in blocks), 
                "z": min(_safe_int(getattr(b, "z_index", 0)) for b in blocks), 
                "x": min(_safe_float(getattr(b, "x", 0.0)) for b in blocks),
                "obj": blocks,
            })

        for item in non_table_items:
            render_queue.append({
                "type": item.kind, "y": item.y, "z": item.z, "x": item.x,
                "obj": item.obj, "w": item.w, "h": item.h, "bottom": item.bottom,
                "alignment": item.alignment, "original_x": item.original_x,
            })

        render_queue.sort(key=lambda e: (e["y"], e["z"], e["x"]))
        rows = ExportService._group_into_rows(render_queue, ROW_GROUP_TOLERANCE_PT)
        last_y_bottom = 0.0

        for row in rows:
            row_top = min(item["y"] for item in row)
            gap = max(0.0, row_top - last_y_bottom)

            has_table = any(item["type"] == "table" for item in row)
            has_manual_table = any(item["type"] == "manual_table" for item in row)

            if has_table:
                for table_item in [i for i in row if i["type"] == "table"]:
                    if gap > 0: _add_spacer_paragraph(doc, gap); gap = 0
                    try: ExportService._add_table_to_docx(doc, table_item["obj"], content_width_pt)
                    except Exception as e: logger.error(f"Failed to render OCR table: {e}")
                    last_y_bottom = max(_safe_float(getattr(b, "y", 0.0)) + _safe_float(getattr(b, "height", 0.0)) for b in table_item["obj"])
            elif has_manual_table:
                for item in [i for i in row if i["type"] == "manual_table"]:
                    if gap > 0: _add_spacer_paragraph(doc, gap); gap = 0
                    try: ExportService._add_manual_table_to_docx(doc, item["obj"], content_width_pt)
                    except Exception as e: logger.error(f"Failed to render manual table: {e}")
                    last_y_bottom = item["y"] + item["h"]
            elif len(row) == 1:
                ExportService._render_single_item(doc, row[0], gap, page_width_pt, content_width_pt)
                last_y_bottom = row[0]["bottom"]
            else:
                ExportService._render_row_as_positioned_table(doc, row, gap, page_width_pt, content_width_pt)
                last_y_bottom = max(item["bottom"] for item in row)

    @staticmethod
    def _collect_layout_items(page: Page, page_width_pt: float) -> List[LayoutItem]:
        blocks = list(page.blocks.all().order_by("y", "x"))
        images = list(PageImage.objects.filter(page=page).order_by("y", "x"))
        manual_tables = list(PageTable.objects.filter(page=page).order_by("y", "x"))

        content_width_pt = page_width_pt - (2 * BASE_MARGIN_PT)
        manual_table_coords = {(round(_safe_float(pt.x), 1), round(_safe_float(pt.y), 1)) for pt in manual_tables}

        items, order = [], 0
        for b in blocks:
            bx, by = _safe_float(getattr(b, "x", 0.0)), _safe_float(getattr(b, "y", 0.0))
            bw, bh = _safe_float(getattr(b, "width", 0.0)), _safe_float(getattr(b, "height", 0.0))
            if getattr(b, "block_type", None) == "table" and (round(bx, 1), round(by, 1)) in manual_table_coords: continue
            
            alignment = _detect_alignment_from_position(bx, bw, page_width_pt, content_width_pt)
            items.append(LayoutItem("block", b, round(bx, 2), round(by, 2), round(bw, 2), round(bh, 2), _safe_int(getattr(b, "z_index", 0)), order, bx, alignment))
            order += 1

        for img in images:
            ix, iy = _safe_float(getattr(img, "x", 0.0)), _safe_float(getattr(img, "y", 0.0))
            iw, ih = _safe_float(getattr(img, "width", 0.0)), _safe_float(getattr(img, "height", 0.0))
            alignment = _detect_alignment_from_position(ix, iw, page_width_pt, content_width_pt)
            items.append(LayoutItem("image", img, round(ix, 2), round(iy, 2), round(iw, 2), round(ih, 2), _safe_int(getattr(img, "z_index", 1)), order, ix, alignment))
            order += 1

        for pt in manual_tables:
            tx, ty = _safe_float(getattr(pt, "x", 0.0)), _safe_float(getattr(pt, "y", 0.0))
            tw, th = _safe_float(getattr(pt, "width", 0.0)), _safe_float(getattr(pt, "height", 0.0))
            items.append(LayoutItem("manual_table", pt, round(tx, 2), round(ty, 2), round(tw, 2), round(th, 2), 0, order, tx, "left"))
            order += 1

        if items:
            min_y = min(item.y for item in items)
            for item in items: item.y = round(item.y - min_y, 2)
            items.sort(key=lambda it: (it.y, it.z, it.x, it.order))
        return items

    @staticmethod
    def _group_into_rows(render_queue: List[Dict[str, Any]], tolerance: float) -> List[List[Dict[str, Any]]]:
        if not render_queue: return []
        rows, current_row, current_y = [], [], None

        for item in render_queue:
            if current_y is None:
                current_row, current_y = [item], item["y"]
            elif abs(item["y"] - current_y) <= tolerance:
                current_row.append(item)
            else:
                if current_row:
                    current_row.sort(key=lambda e: (e["x"], e["z"]))
                    rows.append(current_row)
                current_row, current_y = [item], item["y"]

        if current_row:
            current_row.sort(key=lambda e: (e["x"], e["z"]))
            rows.append(current_row)
        return rows

    @staticmethod
    def _split_tables(items: List[LayoutItem]) -> Tuple[Dict[Any, List[Block]], List[LayoutItem]]:
        table_groups, non_table_items = defaultdict(list), []
        for item in items:
            if item.kind == "block" and getattr(item.obj, "block_type", None) == "table_cell" and getattr(item.obj, "table_id", None):
                table_groups[getattr(item.obj, "table_id")].append(item.obj)
            else:
                non_table_items.append(item)

        for tid in table_groups:
            table_groups[tid].sort(key=lambda b: (
                getattr(b, "row_index", 0) or 0, getattr(b, "col_index", 0) or 0,
                _safe_float(getattr(b, "y", 0.0)), _safe_float(getattr(b, "x", 0.0))
            ))
        return table_groups, non_table_items

    @staticmethod
    def _render_single_item(doc, item: Dict[str, Any], gap: float, page_width_pt: float, content_width_pt: float) -> None:
        if item["type"] == "image":
            if gap > 0: _add_spacer_paragraph(doc, gap)
            ExportService._add_image_with_precise_position(doc, item["obj"], item.get("alignment", "left"), page_width_pt)
            return

        if item["type"] == "block":
            block = item["obj"]
            if getattr(block, "block_type", None) in ("line", "double_line"):
                if gap > 0: _add_spacer_paragraph(doc, gap)
                ExportService._add_graphical_line_as_table(doc, block, 0.0)
                return

            if gap > 0: _add_spacer_paragraph(doc, gap)
            p = doc.add_paragraph()
            ExportService._apply_precise_paragraph_style(p, block, item.get("alignment", "left"), item.get("original_x", item["x"]), item["w"], page_width_pt, content_width_pt)
            
            txt = _clean_text((getattr(block, "current_text", None) or getattr(block, "original_text", None) or "").strip())
            if txt: ExportService._apply_run_style(p.add_run(txt), block)

    @staticmethod
    def _render_row_as_positioned_table(doc, row: List[Dict[str, Any]], gap: float, page_width_pt: float, content_width_pt: float) -> None:
        """Renders multiple items horizontally (like TOC lines). Ensures width doesn't exceed page."""
        if gap > 0: _add_spacer_paragraph(doc, gap)
        if not row: return

        row.sort(key=lambda e: e["x"])
        columns = ExportService._calculate_column_layout(row, content_width_pt)
        if not columns: columns = [{"width": content_width_pt, "item": row[0]}]

        table = doc.add_table(rows=1, cols=len(columns))
        table.autofit = False
        _apply_borderless_style(table)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        _set_table_indent(table, row[0]["x"])

        # WIDTH CAPPING: Prevent Table from overflowing right margin
        total_width = sum(col["width"] for col in columns)
        if total_width > content_width_pt:
            scale = content_width_pt / total_width
            for col in columns: col["width"] *= scale
            total_width = content_width_pt
            
        _set_table_width(table, total_width)

        for idx, col_info in enumerate(columns):
            _set_column_width(table.columns[idx], col_info["width"])
            cell = table.cell(0, idx)
            _set_cell_width(cell, col_info["width"])
            
            tc_pr = cell._tc.get_or_add_tcPr()
            if tc_pr.find(qn("w:vAlign")) is not None: tc_pr.remove(tc_pr.find(qn("w:vAlign")))
            v_al = OxmlElement("w:vAlign")
            v_al.set(qn("w:val"), "top")
            tc_pr.append(v_al)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = LINE_SPACING_DEFAULT

            if col_info["item"] is None: continue

            item = col_info["item"]
            if item["type"] == "image":
                ExportService._render_image_in_cell(p, item["obj"])
                continue

            if item["type"] == "block":
                block = item["obj"]
                txt = _clean_text((getattr(block, "current_text", None) or getattr(block, "original_text", None) or "").strip())
                
                align = item.get("alignment", "left")
                if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                lk = _detect_list_kind(txt)
                p.style = "List Bullet" if lk == "bullet" else "List Number" if lk == "number" else "Normal"
                if txt: ExportService._apply_run_style(p.add_run(txt), block)

    @staticmethod
    def _calculate_column_layout(row: List[Dict[str, Any]], content_width_pt: float) -> List[Dict[str, Any]]:
        if not row: return []
        columns, current_end = [], row[0]["x"]

        for item in row:
            item_start = item["x"]
            item_width = max(MIN_COLUMN_WIDTH_PT, item.get("w", MIN_COLUMN_WIDTH_PT))
            item_end = item_start + item_width
            
            if item_start > current_end + X_POSITION_TOLERANCE_PT:
                columns.append({"width": item_start - current_end, "item": None})
            elif item_start < current_end - X_POSITION_TOLERANCE_PT:
                if columns and columns[-1]["item"] is not None:
                    columns[-1]["width"] = max(MIN_COLUMN_WIDTH_PT, columns[-1]["width"] + (item_end - current_end))
                else:
                    columns.append({"width": item_width, "item": item})
                current_end = max(current_end, item_end)
                continue
            
            columns.append({"width": item_width, "item": item})
            current_end = item_end
        return columns

    @staticmethod
    def _add_image_with_precise_position(doc, img_obj: PageImage, alignment: str, page_width_pt: float) -> None:
        if not getattr(img_obj, "image_file", None): return
        try:
            p = doc.add_paragraph()
            x, w, h = _safe_float(getattr(img_obj, "x", 0.0)), _safe_float(getattr(img_obj, "width", 0.0)), _safe_float(getattr(img_obj, "height", 0.0))

            if alignment == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.left_indent = Pt(0)
            elif alignment == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.left_indent = Pt(0)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                ind = max(0.0, x - BASE_MARGIN_PT)
                p.paragraph_format.left_indent = Pt(0.0 if ind < MARGIN_SNAP_TOLERANCE else ind)

            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0 
            img_obj.image_file.seek(0)
            p.add_run().add_picture(img_obj.image_file, width=Inches(max(0.01, w / 72.0)), height=Inches(max(0.01, h / 72.0)))
        except Exception as e: logger.error(f"Failed to add image: {e}")

    @staticmethod
    def _render_image_in_cell(paragraph, img_obj: PageImage) -> None:
        if not getattr(img_obj, "image_file", None): return
        try:
            img_obj.image_file.seek(0)
            paragraph.add_run().add_picture(img_obj.image_file, width=Inches(max(0.01, _safe_float(getattr(img_obj, "width", 0.0)) / 72.0)), height=Inches(max(0.01, _safe_float(getattr(img_obj, "height", 0.0)) / 72.0)))
        except Exception: pass

    @staticmethod
    def _add_table_to_docx(doc, table_blocks: List[Block], content_width_pt: float) -> None:
        rows = [b.row_index for b in table_blocks if getattr(b, "row_index", None) is not None]
        cols = [b.col_index for b in table_blocks if getattr(b, "col_index", None) is not None]
        if not rows or not cols: return

        max_row, max_col = max(rows), max(cols)
        table = doc.add_table(rows=max_row + 1, cols=max_col + 1)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _apply_borderless_style(table)

        col_left, col_right = {}, {}
        for b in table_blocks:
            col = getattr(b, "col_index", None)
            if col is None: continue
            x, w = _safe_float(getattr(b, "x", 0.0)), _safe_float(getattr(b, "width", 0.0))
            if col not in col_left or x < col_left[col]: col_left[col] = x
            if col not in col_right or (x + w) > col_right[col]: col_right[col] = x + w

        col_widths = {}
        for col in sorted(col_left.keys()):
            next_col = col + 1
            width = col_left[next_col] - col_left[col] if next_col in col_left else col_right[col] - col_left[col]
            col_widths[col] = max(MIN_COLUMN_WIDTH_PT, width)

        # Capping table width to prevent overflow
        total_w = sum(col_widths.values())
        if total_w > content_width_pt:
            scale = content_width_pt / total_w
            for k in col_widths: col_widths[k] *= scale

        for col_idx, width_pt in col_widths.items():
            if 0 <= col_idx <= max_col and width_pt > 0: _set_column_width(table.columns[col_idx], width_pt)

        for b in table_blocks:
            row, col = getattr(b, "row_index", None), getattr(b, "col_index", None)
            if row is None or col is None: continue

            cell = table.cell(row, col)
            rs, cs = _safe_int(getattr(b, "rowspan", 1), 1), _safe_int(getattr(b, "colspan", 1), 1)
            if rs > 1 or cs > 1:
                try: cell.merge(table.cell(min(max_row, row + rs - 1), min(max_col, col + cs - 1)))
                except Exception: pass

            v_align_override = getattr(b, "vertical_align", None)
            val = {"top": "top", "middle": "center", "center": "center", "bottom": "bottom"}.get(str(v_align_override).lower()) if v_align_override else "top"
            
            tc_pr = cell._tc.get_or_add_tcPr()
            if tc_pr.find(qn("w:vAlign")) is not None: tc_pr.remove(tc_pr.find(qn("w:vAlign")))
            v_al = OxmlElement("w:vAlign")
            v_al.set(qn("w:val"), val)
            tc_pr.append(v_al)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = LINE_SPACING_DEFAULT; p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            txt = _clean_text((getattr(b, "current_text", None) or getattr(b, "original_text", None) or "").strip())
            if txt:
                lk = _detect_list_kind(txt)
                p.style = "List Bullet" if lk == "bullet" else "List Number" if lk == "number" else "Normal"
                ExportService._apply_run_style(p.add_run(txt), b)

            role = getattr(b, "semantic_role", "") or ""
            if role in ("heading_1", "heading_2"):
                shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
                cell._tc.get_or_add_tcPr().append(shd)

    @staticmethod
    def _add_manual_table_to_docx(doc, page_table: PageTable, content_width_pt: float) -> None:
        table_json = getattr(page_table, "table_json", None) or []
        has_borders = getattr(page_table, "has_borders", True)
        tbl_width = _safe_float(getattr(page_table, "width", 0.0))
        col_widths_pct: List[float] = getattr(page_table, "col_widths", None) or []

        rows = len(table_json) if table_json else _safe_int(getattr(page_table, "row_count", 0))
        cols = max(len(row) for row in table_json) if table_json else _safe_int(getattr(page_table, "col_count", 0))

        if rows == 0 or cols == 0: return

        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        if has_borders: table.style = "Table Grid"
        else: _apply_borderless_style(table)
        
        # Width capping
        if tbl_width > content_width_pt: tbl_width = content_width_pt

        if col_widths_pct and tbl_width > 0 and len(col_widths_pct) == cols:
            total_pct = sum(col_widths_pct)
            if total_pct != 100.0 and total_pct > 0: col_widths_pct = [p * 100.0 / total_pct for p in col_widths_pct]
            for ci, pct in enumerate(col_widths_pct): _set_column_width(table.columns[ci], max(MIN_COLUMN_WIDTH_PT, tbl_width * pct / 100.0))
        elif tbl_width > 0 and cols > 0:
            for ci in range(cols): _set_column_width(table.columns[ci], max(MIN_COLUMN_WIDTH_PT, tbl_width / cols))

        for ri in range(rows):
            for ci in range(cols):
                cell = table.cell(ri, ci)
                tc_pr = cell._tc.get_or_add_tcPr()
                if tc_pr.find(qn("w:vAlign")) is None:
                    v_al = OxmlElement("w:vAlign")
                    v_al.set(qn("w:val"), "top")
                    tc_pr.append(v_al)

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = LINE_SPACING_DEFAULT; p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                cell_data = None
                try: cell_data = table_json[ri][ci]
                except (IndexError, TypeError): pass
                if cell_data is None: continue

                if isinstance(cell_data, dict):
                    txt = _clean_text(str(cell_data.get("text", "") or ""))
                    if _safe_float(cell_data.get("indent", 0)) > 0: p.paragraph_format.left_indent = Pt(_safe_float(cell_data.get("indent", 0)))
                else: txt = _clean_text(str(cell_data))

                if txt:
                    run = p.add_run(txt)
                    run.font.name = "Calibri"; run.font.size = Pt(11)

    @staticmethod
    def _add_graphical_line_as_table(doc, block: Block, space_before: float) -> None:
        if space_before > 0: _add_spacer_paragraph(doc, space_before)

        line_width = max(1.0, _safe_float(getattr(block, "width", 1.0), 1.0))
        line_height = max(0.5, _safe_float(getattr(block, "height", 0.5), 0.5))

        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        _set_column_width(table.columns[0], line_width)
        table.rows[0].height = Pt(line_height)

        _set_table_indent(table, _safe_float(getattr(block, "x", 0.0)))
        _apply_borderless_style(table)

        cell = table.cell(0, 0)
        tc_pr = cell._tc.get_or_add_tcPr()
        if tc_pr.find(qn("w:tcBorders")) is not None: tc_pr.remove(tc_pr.find(qn("w:tcBorders")))

        tc_borders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}"); b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0"); b.set(qn("w:space"), "0"); b.set(qn("w:color"), "auto")
            tc_borders.append(b)

        block_type = getattr(block, "block_type", None)
        if block_type in ("line", "double_line"):
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "double" if block_type == "double_line" else "single")
            bottom.set(qn("w:sz"), str(max(1, int(line_height * 8))))
            bottom.set(qn("w:space"), "0"); bottom.set(qn("w:color"), "000000")
            tc_borders.append(bottom)
        else:
            bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "none"); bottom.set(qn("w:sz"), "0"); bottom.set(qn("w:space"), "0"); bottom.set(qn("w:color"), "auto")
            tc_borders.append(bottom)

        tc_pr.append(tc_borders)
        p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0

    @staticmethod
    def _apply_precise_paragraph_style(p, block: Block, alignment: str, original_x: float, width: float, page_width_pt: float, content_width_pt: float) -> None:
        p.style = "Normal"
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING_DEFAULT

        if alignment == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.left_indent = Pt(0); p.paragraph_format.right_indent = Pt(0)
        elif alignment == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.left_indent = Pt(0); p.paragraph_format.right_indent = Pt(0)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ind = max(0.0, original_x - BASE_MARGIN_PT)
            p.paragraph_format.left_indent = Pt(0.0 if ind < MARGIN_SNAP_TOLERANCE else ind)
            p.paragraph_format.right_indent = Pt(0)

        role = getattr(block, "semantic_role", "") or ""
        if role.startswith("heading"):
            p.paragraph_format.keep_with_next = True; p.paragraph_format.keep_together = True

    @staticmethod
    def _apply_run_style(run, block: Block) -> None:
        run.font.name = _map_font(getattr(block, "font_name", "") or "")
        run.font.size = Pt(_safe_float(getattr(block, "font_size", 10.0), 10.0))

        role, weight, style = getattr(block, "semantic_role", "") or "", (getattr(block, "font_weight", "") or "").lower(), (getattr(block, "font_style", "") or "").lower()
        run.bold = bool(weight == "bold" or role.startswith("heading"))
        run.italic = bool(style == "italic" or role == "footnote")
        _apply_underline(run, block)

        if getattr(block, "is_strikethrough", False):
            rpr = run._r.get_or_add_rPr()
            strike = OxmlElement("w:strike"); strike.set(qn("w:val"), "1")
            rpr.append(strike)

        spacing = getattr(block, "char_spacing", None)
        if spacing is not None:
            try:
                rpr = run._r.get_or_add_rPr()
                spacing_el = OxmlElement("w:spacing"); spacing_el.set(qn("w:val"), str(int(float(spacing) * 20)))
                rpr.append(spacing_el)
            except Exception: pass

        color = _parse_color(getattr(block, "font_color", "") or "#000000")
        if color: run.font.color.rgb = color