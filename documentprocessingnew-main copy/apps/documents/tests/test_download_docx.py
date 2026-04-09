import os
import shutil
from io import BytesIO
from unittest import mock

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from docx import Document as DocxDocument

from apps.desktop_bridge.models import AssignmentBundle, PageVersion, UploadedPDF
from apps.documents.models import Document, Page
from apps.documents.views import _build_approved_submissions_word_export
from apps.processing.models import PageAssignment, SubmittedPage
from common.enums import PageAssignmentStatus, ReviewStatus


User = get_user_model()
TEST_MEDIA_ROOT = os.path.join(os.getcwd(), "tmp_test_media")
os.makedirs(TEST_MEDIA_ROOT, exist_ok=True)


@override_settings(MEDIA_ROOT=TEST_MEDIA_ROOT)
class DocumentDownloadDocxMergeTest(TestCase):
    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()
        shutil.rmtree(TEST_MEDIA_ROOT, ignore_errors=True)

    def setUp(self):
        self.client_user = User.objects.create_user(
            username="download_client",
            email="download@example.com",
            password="password123",
            role="CLIENT",
        )
        self.document = Document.objects.create(
            client=self.client_user,
            title="Download Merge Test",
            total_pages=2,
        )
        self.resource_user = User.objects.create_user(
            username="download_resource",
            email="resource@example.com",
            password="password123",
            role="RESOURCE",
        )
        self.resource_profile = self.resource_user.resource_profile
        self.resource_profile.is_available = True
        self.resource_profile.status = "ACTIVE"
        self.resource_profile.save(update_fields=["is_available", "status"])
        self.page1 = Page.objects.create(document=self.document, page_number=1)
        self.page2 = Page.objects.create(document=self.document, page_number=2)

    def _docx_upload(self, filename, text):
        buffer = BytesIO()
        doc = DocxDocument()
        doc.add_paragraph(text)
        doc.save(buffer)
        buffer.seek(0)
        return SimpleUploadedFile(
            filename,
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    @mock.patch("apps.processing.tasks.assign_pages_task.delay")
    @mock.patch("apps.processing.tasks.merge_document_pages.delay")
    def test_merge_uses_desktop_docx_when_submission_artifact_is_pdf(self, _mock_merge_delay, _mock_assign_delay):
        assignment1 = PageAssignment.objects.create(
            page=self.page1,
            document=self.document,
            resource=self.resource_profile,
            status=PageAssignmentStatus.SUBMITTED,
        )
        assignment2 = PageAssignment.objects.create(
            page=self.page2,
            document=self.document,
            resource=self.resource_profile,
            status=PageAssignmentStatus.SUBMITTED,
        )

        submission1 = SubmittedPage.objects.create(
            assignment=assignment1,
            page=self.page1,
            document=self.document,
            page_number=1,
            review_status=ReviewStatus.APPROVED,
        )
        submission1.output_page_file.save(
            "page1.pdf",
            SimpleUploadedFile("page1.pdf", b"%PDF-1.4 fake", content_type="application/pdf"),
            save=True,
        )

        submission2 = SubmittedPage.objects.create(
            assignment=assignment2,
            page=self.page2,
            document=self.document,
            page_number=2,
            review_status=ReviewStatus.APPROVED,
        )
        submission2.output_page_file.save(
            "page2.docx",
            self._docx_upload("page2.docx", "Submitted page two"),
            save=True,
        )

        bundle = AssignmentBundle.objects.create(
            document=self.document,
            bundle_index=1,
            page_start=1,
            page_end=1,
            page_numbers=[1],
        )
        upload = UploadedPDF.objects.create(
            document=self.document,
            bundle=bundle,
            file=self._docx_upload("page1.docx", "Desktop page one"),
            checksum="desktop-docx",
        )
        PageVersion.objects.create(
            document=self.document,
            page=self.page1,
            bundle=bundle,
            uploaded_pdf=upload,
            page_number=1,
            page_index_in_pdf=0,
            slice_size=1,
        )

        buffer, filename = _build_approved_submissions_word_export(self.document)

        merged = DocxDocument(buffer)
        merged_text = "\n".join(p.text for p in merged.paragraphs if p.text.strip())

        self.assertTrue(filename.endswith("_merged.docx"))
        self.assertIn("Desktop page one", merged_text)
        self.assertIn("Submitted page two", merged_text)
        self.assertLess(merged_text.index("Desktop page one"), merged_text.index("Submitted page two"))
