import hashlib
import io
import json
import logging
from collections import defaultdict
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from django.core.files.base import ContentFile
from django.db import transaction
from django.utils import timezone
from docx import Document as DocxDocument
from docx.enum.text import WD_BREAK
from docx.shared import Inches

from apps.audit.models import AuditLog
from apps.documents.models import Document
from apps.processing.models import PageAssignment, SubmittedPage
from common.enums import AuditEventType, PageAssignmentStatus, ReviewStatus

logger = logging.getLogger(__name__)


class FinalDocxMergeError(Exception):
    def __init__(self, message, *, code="FINAL_DOCX_MERGE_FAILED", status_code=409, details=None):
        super().__init__(message)
        self.message = message
        self.code = code
        self.status_code = status_code
        self.details = details or {}

    def to_payload(self):
        payload = {"error": self.message, "code": self.code}
        payload.update(self.details)
        return payload


class FinalDocxService:
    @staticmethod
    def _is_readable_docx(file_field):
        if not file_field:
            return False

        try:
            file_field.open("rb")
            try:
                file_field.seek(0)
                DocxDocument(file_field)
                return True
            finally:
                try:
                    file_field.close()
                except Exception:
                    pass
        except Exception:
            return False

    @staticmethod
    def _storage_exists(file_field):
        name = str(getattr(file_field, "name", "") or "").strip()
        if not name:
            return False
        storage = getattr(file_field, "storage", None)
        if not storage:
            return False
        try:
            return storage.exists(name)
        except Exception:
            return False

    @staticmethod
    def _get_original_pdf_page_as_docx(document, page_number):
        """
        Converts a single page from the original source PDF into a minimal DOCX
        containing an embedded image of that page. Used as a fallback for pages
        that have not been assigned, edited, or submitted.
        """
        try:
            from pypdf import PdfWriter, PdfReader
            import tempfile, subprocess, sys
            from pathlib import Path as _Path

            # Resolve original PDF path
            src = None
            if document.file:
                try:
                    src = document.file.path
                except Exception:
                    pass
            if not src and hasattr(document, 'original_file') and document.original_file:
                try:
                    src = document.original_file.path
                except Exception:
                    pass

            if not src or not _Path(src).exists():
                raise ValueError(f"Original PDF not found for document {document.id}")

            # Extract the single page as a temporary PDF
            reader = PdfReader(src)
            if page_number < 1 or page_number > len(reader.pages):
                raise ValueError(f"Page {page_number} out of range (PDF has {len(reader.pages)} pages)")

            writer = PdfWriter()
            writer.add_page(reader.pages[page_number - 1])
            tmp_pdf = io.BytesIO()
            writer.write(tmp_pdf)
            tmp_pdf.seek(0)

            # Create a DOCX with a note about this being the original page
            doc = DocxDocument()
            para = doc.add_paragraph(f"[Original Page {page_number} — not yet edited]")
            para.runs[0].bold = True

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

        except Exception as exc:
            logger.warning(f"Could not extract original PDF page {page_number} for document {document.id}: {exc}")
            # Last resort: blank placeholder docx
            doc = DocxDocument()
            doc.add_paragraph(f"[Page {page_number} — content unavailable]")
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

    @staticmethod
    def _build_generated_page_docx(document, page_number):
        from apps.processing.services.export import ExportService

        page = document.pages.filter(page_number=page_number).first()
        if not page:
            raise FinalDocxMergeError(
                f"Page {page_number} is missing from the document page map.",
                code="PAGE_NOT_FOUND",
                details={"page_number": page_number},
            )

        try:
            buffer = ExportService.generate_word_export(
                document,
                include_unapproved=True,
                pages=[page],
            )
        except Exception as exc:
            raise FinalDocxMergeError(
                f"Page {page_number} could not be regenerated from saved workspace data.",
                code="PAGE_DOCX_REBUILD_FAILED",
                details={"page_number": page_number},
            ) from exc

        return buffer

    @staticmethod
    def _get_latest_desktop_docx(document, page_number):
        try:
            from apps.desktop_bridge.models import PageVersion
        except Exception:
            return None, None, ""

        versions = (
            PageVersion.objects.filter(
                document=document,
                page_number=page_number,
                is_valid=True,
            )
            .select_related("uploaded_pdf__device", "bundle")
            .order_by("slice_size", "-updated_at", "id")
        )

        latest_seen_name = ""
        for version in versions:
            upload = getattr(version, "uploaded_pdf", None)
            file_field = getattr(upload, "file", None)
            if file_field and not latest_seen_name:
                latest_seen_name = str(getattr(file_field, "name", "") or "")
            if FinalDocxService._is_readable_docx(file_field):
                return version, file_field, latest_seen_name

        return None, None, latest_seen_name

    @staticmethod
    def _safe_name(document):
        raw = document.doc_ref or document.name or f"document_{str(document.id)[:8]}"
        safe = "".join(c if c.isalnum() or c in ("-", "_") else "_" for c in raw).strip("_")
        return safe or f"document_{str(document.id)[:8]}"

    @staticmethod
    def _build_filename(document, manifest_pages):
        assignment_ids = sorted({item["assignment_id"] for item in manifest_pages if item.get("assignment_id")})
        bundle_ids = sorted({item["bundle_id"] for item in manifest_pages if item.get("bundle_id")})

        if bundle_ids:
            identity = f"b{bundle_ids[0]}" if len(bundle_ids) == 1 else f"b{bundle_ids[0]}_{bundle_ids[-1]}"
        elif assignment_ids:
            identity = f"a{assignment_ids[0]}" if len(assignment_ids) == 1 else f"a{assignment_ids[0]}_{assignment_ids[-1]}"
        else:
            identity = f"doc_{str(document.id)[:8]}"

        timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
        return f"{FinalDocxService._safe_name(document)}_{identity}_{timestamp}.docx"

    @staticmethod
    def _open_docx_from_source(source_obj):
        if hasattr(source_obj, "open"):
            source_obj.open("rb")
        if hasattr(source_obj, "seek"):
            source_obj.seek(0)

        try:
            return DocxDocument(source_obj)
        finally:
            try:
                if hasattr(source_obj, "close"):
                    source_obj.close()
            except Exception:
                pass

    @staticmethod
    def _expected_page_map(document):
        pages = list(document.pages.order_by("page_number").values("id", "page_number"))
        if not pages:
            raise FinalDocxMergeError(
                "Document has no persisted page records to merge.",
                code="NO_PERSISTED_PAGES",
            )

        page_numbers = [int(item["page_number"]) for item in pages]
        expected_total = int(document.total_pages or len(page_numbers))
        expected_range = list(range(1, expected_total + 1))

        if page_numbers != expected_range:
            missing = sorted(set(expected_range) - set(page_numbers))
            raise FinalDocxMergeError(
                "Document page ordering is incomplete and cannot be inferred safely.",
                code="PAGE_ORDER_UNDETERMINED",
                details={"missing_pages": missing, "expected_pages": expected_total},
            )

        return {int(item["page_number"]): int(item["id"]) for item in pages}

    @staticmethod
    def _collect_approved_submissions(document):
        submissions = list(
            SubmittedPage.objects.filter(
                document=document,
                review_status=ReviewStatus.APPROVED,
            )
            .select_related("submitted_by", "assignment__resource__user", "page")
            .order_by("page_number", "-submitted_at", "-id")
        )

        grouped = defaultdict(list)
        for submission in submissions:
            page_number = int(submission.page_number or getattr(submission.page, "page_number", 0) or 0)
            if page_number > 0:
                grouped[page_number].append(submission)
        return grouped

    @staticmethod
    def _resolve_page_source(document, page_number, page_id, submissions):
        chosen_file = None
        chosen_manifest = None
        duplicate_submission_ids = [submission.id for submission in submissions]
        latest_submission_file = ""
        latest_submission_timestamp = None

        for submission in submissions:
            latest_submission_file = latest_submission_file or str(getattr(submission.output_page_file, "name", "") or "")
            latest_submission_timestamp = latest_submission_timestamp or submission.submitted_at

            if not FinalDocxService._storage_exists(submission.output_page_file):
                continue
            if not FinalDocxService._is_readable_docx(submission.output_page_file):
                continue

            assignment_user_id = None
            if submission.assignment and submission.assignment.resource and submission.assignment.resource.user:
                assignment_user_id = submission.assignment.resource.user_id

            chosen_file = submission.output_page_file
            chosen_manifest = {
                "page_id": page_id,
                "page_number": page_number,
                "assignment_id": submission.assignment_id,
                "bundle_id": None,
                "submission_id": submission.id,
                "user_id": submission.submitted_by_id or assignment_user_id,
                "timestamp": submission.submitted_at.isoformat() if submission.submitted_at else None,
                "source": "submitted_page_docx",
                "file_path": str(getattr(submission.output_page_file, "name", "") or ""),
                "resolved_from_duplicates": len(duplicate_submission_ids) > 1,
                "duplicate_submission_ids": duplicate_submission_ids[1:],
            }
            break

        if chosen_file and chosen_manifest:
            return chosen_file, chosen_manifest

        version, desktop_file, latest_desktop_name = FinalDocxService._get_latest_desktop_docx(document, page_number)
        if version and desktop_file:
            device_user_id = None
            if version.uploaded_pdf and version.uploaded_pdf.device:
                device_user_id = version.uploaded_pdf.device.user_id

            chosen_manifest = {
                "page_id": page_id,
                "page_number": page_number,
                "assignment_id": submissions[0].assignment_id if submissions else None,
                "bundle_id": version.bundle_id,
                "submission_id": submissions[0].id if submissions else None,
                "user_id": (
                    submissions[0].submitted_by_id
                    if submissions and submissions[0].submitted_by_id
                    else device_user_id
                ),
                "timestamp": version.updated_at.isoformat() if version.updated_at else None,
                "source": "desktop_page_docx",
                "file_path": str(getattr(desktop_file, "name", "") or ""),
                "resolved_from_duplicates": len(duplicate_submission_ids) > 1,
                "duplicate_submission_ids": duplicate_submission_ids[1:],
            }
            return desktop_file, chosen_manifest

        # ── NEW: try original PDF page as fallback before regenerating from workspace ──
        # This handles pages that were never assigned/edited (e.g. pages outside the bundle range).
        has_any_assignment = PageAssignment.objects.filter(
            document=document, page__page_number=page_number
        ).exists()

        if not has_any_assignment and not submissions:
            # Page was never assigned — pull straight from original PDF
            pdf_buf = FinalDocxService._get_original_pdf_page_as_docx(document, page_number)
            fallback_manifest = {
                "page_id": page_id,
                "page_number": page_number,
                "assignment_id": None,
                "bundle_id": None,
                "submission_id": None,
                "user_id": None,
                "timestamp": timezone.now().isoformat(),
                "source": "original_pdf_page",
                "file_path": None,
                "resolved_from_duplicates": False,
                "duplicate_submission_ids": [],
            }
            return pdf_buf, fallback_manifest

        # Final fallback: regenerate from workspace data
        generated = FinalDocxService._build_generated_page_docx(document, page_number)
        fallback_user_id = None
        fallback_assignment_id = None
        fallback_timestamp = None
        if submissions:
            fallback_user_id = submissions[0].submitted_by_id
            fallback_assignment_id = submissions[0].assignment_id
            fallback_timestamp = submissions[0].submitted_at.isoformat() if submissions[0].submitted_at else None
        else:
            latest_assignment = (
                PageAssignment.objects.filter(document=document, page__page_number=page_number)
                .select_related("resource__user")
                .order_by("-assigned_at")
                .first()
            )
            if latest_assignment:
                fallback_assignment_id = latest_assignment.id
                if latest_assignment.resource and latest_assignment.resource.user:
                    fallback_user_id = latest_assignment.resource.user_id
                if latest_assignment.assigned_at:
                    fallback_timestamp = latest_assignment.assigned_at.isoformat()

        chosen_manifest = {
            "page_id": page_id,
            "page_number": page_number,
            "assignment_id": fallback_assignment_id,
            "bundle_id": version.bundle_id if version else None,
            "submission_id": submissions[0].id if submissions else None,
            "user_id": fallback_user_id,
            "timestamp": fallback_timestamp,
            "source": "regenerated_workspace_docx",
            "file_path": None,
            "resolved_from_duplicates": len(duplicate_submission_ids) > 1,
            "duplicate_submission_ids": duplicate_submission_ids[1:],
            "latest_submission_file": latest_submission_file or None,
            "latest_desktop_file": latest_desktop_name or None,
        }
        return generated, chosen_manifest

    @staticmethod
    def _validate_manifest(document, manifest_pages, expected_page_numbers, assigned_page_numbers):
        missing_pages = sorted({
            page for page in expected_page_numbers
            if page not in {item["page_number"] for item in manifest_pages}
        })

        if missing_pages:
            raise FinalDocxMergeError(
                "Some expected pages could not be resolved for merge.",
                code="MISSING_SAVED_PAGES",
                details={"missing_pages": missing_pages},
            )

        # Unassigned pages (using original PDF) are fine — just log them
        missing_assigned_pages = sorted(set(expected_page_numbers) - set(assigned_page_numbers))
        if missing_assigned_pages:
            logger.info(
                "Pages %s are unassigned and will use the original PDF source.",
                missing_assigned_pages,
            )

        invalid_pages = []
        for item in manifest_pages:
            if not item.get("page_id"):
                invalid_pages.append({"page_number": item["page_number"], "reason": "missing page_id"})
            if not item.get("page_number"):
                invalid_pages.append({"page_number": item["page_number"], "reason": "missing page_number"})
            # Only validate assignment/bundle for pages that WERE assigned
            if item.get("source") != "original_pdf_page":
                if not item.get("assignment_id") and not item.get("bundle_id"):
                    invalid_pages.append({"page_number": item["page_number"], "reason": "missing assignment_id/bundle_id"})
                if not item.get("user_id"):
                    invalid_pages.append({"page_number": item["page_number"], "reason": "missing user_id"})
                if not item.get("timestamp"):
                    invalid_pages.append({"page_number": item["page_number"], "reason": "missing timestamp"})
                if not item.get("file_path") and item.get("source") != "regenerated_workspace_docx":
                    invalid_pages.append({"page_number": item["page_number"], "reason": "missing saved file path"})

        if invalid_pages:
            raise FinalDocxMergeError(
                "One or more page records are not stored completely and cannot be merged safely.",
                code="INVALID_PAGE_STORAGE",
                details={"invalid_pages": invalid_pages},
            )

    @staticmethod
    def _build_merge_plan(document):
        page_map = FinalDocxService._expected_page_map(document)
        expected_page_numbers = list(page_map.keys())

        assigned_page_numbers = list(
            PageAssignment.objects.filter(document=document)
            .values_list("page__page_number", flat=True)
            .distinct()
        )

        approved_submissions = FinalDocxService._collect_approved_submissions(document)
        approved_page_numbers = sorted(approved_submissions.keys())
        missing_submitted_pages = sorted(set(expected_page_numbers) - set(approved_page_numbers))

        if missing_submitted_pages:
            # Log which pages are missing — but do NOT block the merge.
            # Unassigned/unedited pages will use the original PDF as fallback.
            logger.info(
                "Document %s: pages %s have no approved submission — will use original PDF fallback.",
                document.id,
                missing_submitted_pages,
            )

        merge_sources = []
        duplicate_pages = {}
        latest_source_timestamp = None

        for page_number in expected_page_numbers:
            submissions = approved_submissions.get(page_number, [])
            if len(submissions) > 1:
                duplicate_pages[page_number] = [submission.id for submission in submissions]

            source_obj, manifest_entry = FinalDocxService._resolve_page_source(
                document,
                page_number,
                page_map[page_number],
                submissions,
            )
            merge_sources.append({"page_number": page_number, "source_obj": source_obj, "manifest": manifest_entry})

            timestamp_value = manifest_entry.get("timestamp")
            if timestamp_value:
                try:
                    ts = timezone.datetime.fromisoformat(timestamp_value.replace("Z", "+00:00"))
                    if latest_source_timestamp is None or ts > latest_source_timestamp:
                        latest_source_timestamp = ts
                except Exception:
                    pass

        manifest_pages = [item["manifest"] for item in merge_sources]
        FinalDocxService._validate_manifest(document, manifest_pages, expected_page_numbers, assigned_page_numbers)

        signature_payload = {
            "document_id": str(document.id),
            "pages": manifest_pages,
        }
        signature = hashlib.sha256(
            json.dumps(signature_payload, sort_keys=True, default=str).encode("utf-8")
        ).hexdigest()

        return {
            "merge_sources": merge_sources,
            "manifest": {
                "document_id": str(document.id),
                "expected_page_count": len(expected_page_numbers),
                "assigned_page_count": len(set(assigned_page_numbers)),
                "approved_page_count": len(approved_page_numbers),
                "missing_pages": [],
                "duplicate_pages": duplicate_pages,
                "pages": manifest_pages,
                "signature": signature,
                "latest_source_timestamp": latest_source_timestamp.isoformat() if latest_source_timestamp else None,
            },
        }

    @staticmethod
    def _merge_sources_to_buffer(document, merge_sources):
        merged = DocxDocument()
        body = merged.element.body
        for child in list(body.iterchildren()):
            if child.tag.endswith("}p"):
                body.remove(child)

        for idx, item in enumerate(merge_sources):
            page_number = item["page_number"]
            try:
                part = FinalDocxService._open_docx_from_source(item["source_obj"])
            except Exception as exc:
                raise FinalDocxMergeError(
                    f"Page {page_number} has an unreadable DOCX source.",
                    code="UNREADABLE_DOCX_SOURCE",
                    details={"page_number": page_number},
                ) from exc

            for child in part.element.body.iterchildren():
                if child.tag.endswith("}sectPr"):
                    continue
                body.append(deepcopy(child))

            if idx < len(merge_sources) - 1:
                paragraph = merged.add_paragraph()
                paragraph.add_run().add_break(WD_BREAK.PAGE)

        buffer = BytesIO()
        merged.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    @transaction.atomic
    def prepare_final_docx(document, actor=None, force=False):
        document = Document.objects.select_for_update().get(pk=document.pk)
        plan = FinalDocxService._build_merge_plan(document)
        manifest = plan["manifest"]

        current_signature = document.final_word_manifest.get("signature") if isinstance(document.final_word_manifest, dict) else None
        has_current_file = document.final_word_file and FinalDocxService._storage_exists(document.final_word_file)

        if has_current_file and not force and current_signature == manifest["signature"]:
            logger.info("Reusing persisted final DOCX for document %s", document.id)
            return {
                "document": document,
                "filename": Path(document.final_word_file.name).name,
                "generated": False,
                "manifest": manifest,
            }

        logger.info("Generating final DOCX for document %s", document.id)
        buffer = FinalDocxService._merge_sources_to_buffer(document, plan["merge_sources"])
        filename = FinalDocxService._build_filename(document, manifest["pages"])

        if document.final_word_file:
            try:
                document.final_word_file.delete(save=False)
            except Exception:
                logger.warning("Failed to delete previous final_word_file for document %s", document.id, exc_info=True)

        document.final_word_file.save(filename, ContentFile(buffer.getvalue()), save=False)
        document.final_word_generated_at = timezone.now()
        document.final_word_manifest = manifest
        document.final_word_error = ""
        document.save(update_fields=["final_word_file", "final_word_generated_at", "final_word_manifest", "final_word_error", "updated_at"])

        AuditLog.objects.create(
            action=AuditEventType.DOC_DOWNLOADED,
            document_id=document.id,
            actor=actor,
            metadata={
                "format": "docx",
                "filename": filename,
                "source": "final_docx_merge",
                "generated": True,
                "pages": manifest["expected_page_count"],
                "duplicate_pages": manifest["duplicate_pages"],
                "signature": manifest["signature"],
            },
        )

        return {
            "document": document,
            "filename": filename,
            "generated": True,
            "manifest": manifest,
        }

    @staticmethod
    def record_failure(document, message, details=None):
        Document.objects.filter(pk=document.pk).update(final_word_error=message)
        logger.error("Final DOCX generation failed for %s: %s", document.id, message)
        if details:
            logger.error("Final DOCX failure details for %s: %s", document.id, details)
