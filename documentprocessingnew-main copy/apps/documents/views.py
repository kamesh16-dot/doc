
from rest_framework import viewsets, permissions, status, parsers, views
from rest_framework.response import Response
from rest_framework.decorators import action
from django.shortcuts import get_object_or_404
from apps.documents.models import Document, Page, Block, BlockEdit
from apps.documents.serializers import DocumentSerializer, DocumentUploadSerializer, PageSerializer, BlockSerializer
from apps.documents.services import DocumentService
from common.enums import UserRole
from common.validators import StatusTransitionValidator
from django.utils import timezone
from django.http import HttpResponse
import logging
logger = logging.getLogger(__name__)


def _is_docx_submission_file(file_field):
    """
    True when the stored artifact is a DOCX file.
    The filename alone is not enough in this project because some rows
    have a .docx suffix while the stored bytes are actually PDF/corrupt.
    """
    if not file_field:
        return False

    try:
        from docx import Document as DocxDocument
        file_field.open("rb")
        try:
            file_field.seek(0)
            DocxDocument(file_field)
            return True
        finally:
            try:
                file_field.close()
            except Exception:
                pass
    except Exception:
        return False


def _build_generated_page_word_export(document, page_number):
    """
    Regenerate a single-page DOCX from the saved workspace state.
    This is the repair path when the uploaded artifact is missing/corrupt.
    """
    from apps.processing.services.export import ExportService

    page = document.pages.filter(page_number=page_number).first()
    if not page:
        raise ValueError(f"Page {page_number} does not exist in the document.")

    buffer = ExportService.generate_word_export(
        document,
        include_unapproved=True,
        pages=[page],
    )
    return buffer


def _get_latest_desktop_docx_upload(document, page_number):
    """
    Resolve the latest desktop-agent upload for a page that is readable as DOCX.
    This is the backup source of truth when SubmittedPage still points to a PDF
    snapshot but the desktop bridge already stored the real edited Word file.
    """
    try:
        from apps.desktop_bridge.models import PageVersion
    except Exception:
        return None, ""

    latest_seen_name = ""
    versions = PageVersion.objects.filter(
        document=document,
        page_number=page_number,
        is_valid=True,
    ).select_related("uploaded_pdf").order_by("slice_size", "-updated_at", "id")

    for version in versions:
        upload = getattr(version, "uploaded_pdf", None)
        file_field = getattr(upload, "file", None)
        if file_field and not latest_seen_name:
            latest_seen_name = str(getattr(file_field, "name", "") or "")
        if _is_docx_submission_file(file_field):
            return file_field, latest_seen_name

    return None, latest_seen_name


def _build_approved_submissions_word_export(document):
    """
    Merge approved per-page DOCX submissions in strict page order (1..N).
    This is used by the Download action so admins get exactly what users saved.
    """
    from io import BytesIO
    from copy import deepcopy
    from docx import Document as DocxDocument
    from docx.enum.text import WD_BREAK
    from apps.processing.models import SubmittedPage
    from common.enums import ReviewStatus

    total_pages = int(document.total_pages or 0)
    if total_pages <= 0:
        raise ValueError("Document has no pages to export.")

    approved_qs = SubmittedPage.objects.filter(
        document=document,
        review_status=ReviewStatus.APPROVED
    ).order_by("page_number", "-submitted_at", "-id")

    approved_all = list(approved_qs)
    approved_by_page = {}
    for sub in approved_all:
        page_num = int(sub.page_number or 0)
        if page_num <= 0:
            continue

        if page_num not in approved_by_page:
            approved_by_page[page_num] = {
                "latest": sub,
                "latest_docx": sub if _is_docx_submission_file(sub.output_page_file) else None,
            }
            continue

        # Queryset is already newest-first per page; keep first "latest".
        if approved_by_page[page_num]["latest_docx"] is None and _is_docx_submission_file(sub.output_page_file):
            approved_by_page[page_num]["latest_docx"] = sub

    if len(approved_by_page) != total_pages:
        found = set(approved_by_page.keys())
        missing = [p for p in range(1, total_pages + 1) if p not in found]
        raise ValueError(
            f"Cannot download Word merge: {len(found)}/{total_pages} approved pages. Missing: {missing}"
        )

    approved = []
    for page_num in range(1, total_pages + 1):
        selected = approved_by_page[page_num].get("latest_docx")
        selected_file = getattr(selected, "output_page_file", None) if selected else None
        desktop_file, latest_desktop_name = _get_latest_desktop_docx_upload(document, page_num)
        if not selected_file and desktop_file:
            selected_file = desktop_file
        generated_buffer = None

        if not selected_file:
            try:
                generated_buffer = _build_generated_page_word_export(document, page_num)
                selected_file = generated_buffer
            except Exception:
                generated_buffer = None

        if not selected_file:
            latest_sub = approved_by_page[page_num].get("latest")
            latest_name = str(getattr(getattr(latest_sub, "output_page_file", None), "name", "") or "")
            desktop_hint = f" Desktop latest file is '{latest_desktop_name}'." if latest_desktop_name else ""
            raise ValueError(
                f"Page {page_num} has no approved Word (.docx) submission. "
                f"Latest approved file is '{latest_name or 'missing'}'.{desktop_hint}"
            )
        approved.append({
            "page_number": page_num,
            "file_field": selected_file,
            "generated_buffer": generated_buffer,
        })

    merged = DocxDocument()
    body = merged.element.body
    # Remove the default empty paragraph; keep section properties.
    for child in list(body.iterchildren()):
        if child.tag.endswith("}p"):
            body.remove(child)

    for idx, approved_page in enumerate(approved):
        file_field = approved_page["file_field"]
        page_num = approved_page["page_number"]
        part = None
        try:
            if hasattr(file_field, "open"):
                file_field.open("rb")
            if hasattr(file_field, "seek"):
                file_field.seek(0)
            part = DocxDocument(file_field)
        except Exception as e:
            raise ValueError(f"Page {page_num} has an unreadable Word submission.") from e
        finally:
            try:
                if hasattr(file_field, "close"):
                    file_field.close()
            except Exception:
                pass

        for child in part.element.body.iterchildren():
            # Avoid duplicating section-properties per part.
            if child.tag.endswith("}sectPr"):
                continue
            body.append(deepcopy(child))

        if idx < len(approved) - 1:
            p = merged.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)

    buffer = BytesIO()
    merged.save(buffer)
    buffer.seek(0)

    safe_name = "".join(
        c if c.isalnum() or c in ("-", "_") else "_"
        for c in (document.doc_ref or document.name or "document")
    ).strip("_") or f"document_{str(document.id)[:8]}"
    filename = f"{safe_name}_merged.docx"
    return buffer, filename


class DocumentViewSet(viewsets.ModelViewSet):
    queryset = Document.objects.all()
    serializer_class = DocumentSerializer
    parser_classes = (parsers.MultiPartParser, parsers.FormParser, parsers.JSONParser)

    def get_queryset(self):
        user = self.request.user
        queryset = Document.active.all()
        
        if user.role == 'CLIENT':
            queryset = queryset.filter(client=user)
        elif user.role == 'RESOURCE' and not (user.is_staff or user.is_superuser):
            from apps.processing.models import PageAssignment
            assigned_doc_ids = PageAssignment.objects.filter(
                resource__user=user
            ).values_list('document_id', flat=True).distinct()
            queryset = queryset.filter(id__in=assigned_doc_ids)
        
        # Filtering
        search = self.request.query_params.get('search')
        if search:
            queryset = queryset.filter(name__icontains=search)
            
        status_filter = self.request.query_params.get('status')
        if status_filter and status_filter != 'ALL':
            from django.db.models import Q
            # Map frontend labels to backend status/pipeline_status values
            if status_filter == 'Ready for Merging':
                queryset = queryset.filter(pipeline_status='ALL_SUBMITTED')
            elif status_filter == 'Completed':
                queryset = queryset.filter(Q(status='COMPLETED') | Q(pipeline_status__in=['MERGED', 'APPROVED']))
            elif status_filter == 'Reviewing':
                queryset = queryset.filter(Q(status='REVIEWING') | Q(pipeline_status='VALIDATING'))
            elif status_filter == 'Started':
                queryset = queryset.filter(pipeline_status='IN_PROGRESS')
            elif status_filter == 'Assigned':
                queryset = queryset.filter(status='ASSIGNED')
            else:
                # Fallback for exact enum key match
                queryset = queryset.filter(Q(status=status_filter) | Q(pipeline_status=status_filter))
            
        return queryset.order_by('-id')

    def destroy(self, request, *args, **kwargs):
        """
        Hard delete implementation with thorough cascading cleanup.
        Satisfies user requirement: "archive/delete from assigned resource, database, and queue".
        """
        from apps.processing.models import PageAssignment, DocumentQueue, SubmittedPage
        from apps.audit.models import AuditLog
        import os
        from django.conf import settings
        
        try:
            document = self.get_object()
            doc_id = document.id
            doc_title = document.title or document.name
            
            # 1. Broad Manual Cleanup for un-linked or loosely linked items
            AuditLog.objects.filter(document_id=doc_id).delete()
            
            # 2. Physical File Cleanup from disk
            def safe_delete_file(file_obj):
                if file_obj and hasattr(file_obj, 'name') and file_obj.name:
                    try:
                        if os.path.exists(file_obj.path):
                            os.remove(file_obj.path)
                    except Exception as e:
                        logger.error(f"Failed to delete file: {e}")

            # Delete all page-level split files
            for p in document.pages.all():
                safe_delete_file(p.content_file)
                
            # Delete submitted page blobs
            for sp in SubmittedPage.objects.filter(document=document):
                safe_delete_file(sp.output_page_file)
                
            # Delete merged document if exists
            from apps.processing.models import MergedDocument
            try:
                md = MergedDocument.objects.get(document=document)
                safe_delete_file(md.merged_file)
            except MergedDocument.DoesNotExist:
                pass

            # Delete document originals and processing files
            safe_delete_file(document.original_file)
            safe_delete_file(document.file)
            safe_delete_file(document.converted_pdf)
            safe_delete_file(document.final_file)

            # 3. Hard Delete the document itself (cascades database-level objects)
            document.delete()
            
            return Response({
                'status': 'deleted', 
                'message': f"Document '{doc_title}' and all associated assignments, queue entries, and audit logs have been permanently removed."
            }, status=status.HTTP_200_OK)
            
        except Exception as e:
            logger.error(f"Deletion failed for document: {str(e)}", exc_info=True)
            return Response(
                {"error": str(e), "detail": "Internal server error during deletion."}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


    def get_serializer_class(self):
        if self.action == 'create':
            return DocumentUploadSerializer
        return DocumentSerializer

    def create(self, request, *args, **kwargs):
        try:
            serializer = self.get_serializer(data=request.data)
            serializer.is_valid(raise_exception=True)
            
            uploaded_file = serializer.validated_data['file']
            filename = getattr(uploaded_file, 'name', '').lower()
            
            # ── GUARD: Prevent Duplicate Bundle Results ──────────
            # If a watcher mis-identifies a bundle result as a new document,
            # we reject it here to keep the dashboard clean.
            if filename.startswith("bundle") and filename.endswith(".docx"):
                logger.warning(f"Rejected attempt to create a new Document from a bundle result: {filename}")
                return Response(
                    {"error": "Bundle result files cannot be imported as new documents. They must be uploaded via the bundle result endpoint."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Create document via service
            doc = DocumentService.create_document(request.user, uploaded_file)
            
            # Trigger processing task
            import os
            ext = os.path.splitext(doc.original_file.name)[1].lower()
            
            # Trigger processing task in background to keep upload response fast
            from common.utils import run_task_background
            if ext in ['.docx', '.doc']:
                from apps.documents.tasks import convert_word_to_pdf
                run_task_background(lambda: convert_word_to_pdf.delay(doc.id))
            else:
                from apps.documents.tasks import split_document_task
                run_task_background(lambda: split_document_task.delay(doc.id))
            
            headers = self.get_success_headers(serializer.data)
            # Pass context for absolute URLs (crucial for nested FileFields)
            # Use doc from DB to ensure all fields like doc_ref are populated
            doc.refresh_from_db()
            response_serializer = DocumentSerializer(doc, context=self.get_serializer_context())
            return Response(response_serializer.data, status=status.HTTP_201_CREATED, headers=headers)
        except Exception as e:
            logger.error(f"Upload failed: {str(e)}", exc_info=True)
            return Response(
                {"error": str(e), "detail": "Internal server error during upload."}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    @action(detail=True, methods=['get'], url_path='unassigned-pages')
    def get_unassigned_pages(self, request, pk=None):
        doc = self.get_object()
        from apps.processing.models import PageAssignment
        from common.enums import PageAssignmentStatus
        
        # Get all page numbers
        all_pages = set(doc.pages.values_list('page_number', flat=True))
        
        # Get assigned page numbers
        assigned_pages = set(PageAssignment.objects.filter(
            document=doc,
            status__in=[PageAssignmentStatus.ASSIGNED, PageAssignmentStatus.IN_PROGRESS, PageAssignmentStatus.SUBMITTED]
        ).values_list('page__page_number', flat=True))
        
        unassigned = sorted(list(all_pages - assigned_pages))
        
        return Response({
            'doc_ref': doc.doc_ref,
            'total_pages': doc.total_pages,
            'unassigned_count': len(unassigned),
            'unassigned_pages': unassigned
        })

    @action(detail=True, methods=['get'], url_path='bundles')
    def get_bundles(self, request, pk=None):
        doc = self.get_object()
        from apps.desktop_bridge.models import AssignmentBundle
        
        bundles = AssignmentBundle.objects.filter(document=doc).order_by('page_start')
        
        bundle_data = []
        for b in bundles:
            file_url = None
            if b.output_file:
                file_url = request.build_absolute_uri(b.output_file.url)
            elif b.result_pdf:
                file_url = request.build_absolute_uri(b.result_pdf.url)
            
            bundle_data.append({
                'bundle_id': str(b.id),
                'bundle_index': b.bundle_index,
                'page_start': b.page_start,
                'page_end': b.page_end,
                'status': b.status,
                'file_location': file_url,
                'uploaded_at': b.uploaded_at,
                'updated_at': b.updated_at
            })
            
        return Response({
            'document_id': str(doc.id),
            'doc_ref': doc.doc_ref,
            'total_bundles': len(bundle_data),
            'bundles': bundle_data
        })

    @action(detail=True, methods=['get'], url_path='progress')
    def get_progress(self, request, pk=None):
        doc = self.get_object()
        from apps.processing.models import PageAssignment, SubmittedPage
        from common.enums import PageAssignmentStatus, ReviewStatus
        
        total = doc.total_pages or 0
        if total == 0:
            return Response({'progress': 0, 'status': doc.pipeline_status})
            
        # Approved is 100% complete for that page
        approved = SubmittedPage.objects.filter(document=doc, review_status=ReviewStatus.APPROVED).count()
        
        # Submitted is 80% complete 
        submitted = SubmittedPage.objects.filter(document=doc, review_status=ReviewStatus.PENDING_REVIEW).count()
        
        # In Progress is 40% complete
        processing = PageAssignment.objects.filter(document=doc, status=PageAssignmentStatus.IN_PROGRESS).count()
        
        # Assigned is 10% complete
        assigned = PageAssignment.objects.filter(document=doc, status=PageAssignmentStatus.ASSIGNED).count()
        
        # Weighted score
        score = (approved * 1.0) + (submitted * 0.8) + (processing * 0.4) + (assigned * 0.1)
        percentage = round((score / total) * 100, 1)
        
        return Response({
            'doc_ref': doc.doc_ref,
            'pipeline_status': doc.pipeline_status,
            'completion_percentage': min(100.0, percentage),
            'pages_approved': approved,
            'pages_submitted_pending_review': submitted,
            'pages_processing': processing,
            'pages_assigned_not_started': assigned,
            'total_pages': total
        })

    @action(detail=True, methods=['post'])
    def retry(self, request, pk=None):
        doc = self.get_object()
        from common.enums import DocumentStatus, PipelineStatus
        
        # Reset bits
        doc.status = DocumentStatus.UPLOADED
        doc.pipeline_status = PipelineStatus.UPLOADED
        doc.pipeline_error = ""
        doc.conversion_error = ""
        doc.save()
        
        # Cleanup old assignments for this document
        from apps.processing.models import PageAssignment
        PageAssignment.objects.filter(document=doc).delete()
        
        # Logic from create() to trigger processing
        import os
        ext = os.path.splitext(doc.original_file.name)[1].lower()
        
        if ext in ['.docx', '.doc']:
            from apps.documents.tasks import convert_word_to_pdf
            convert_word_to_pdf.delay(doc.id)
        else:
            # For PDF, if it already had pages, split_document service handles idempotency
            # but we want to force re-processing if it failed.
            # If doc.file is missing but it's a PDF, we might need to restore it from original_file
            if not doc.file and ext == '.pdf':
                doc.file = doc.original_file
                doc.save()
                
            from apps.documents.tasks import split_document_task
            split_document_task.delay(doc.id)
            
        return Response({'status': 'processing_restarted'})

    @action(detail=True, methods=['get'], url_path='assign-all')
    def assign_all(self, request, pk=None):
        doc = self.get_object()
        from common.enums import PipelineStatus
        
        if doc.pipeline_status not in [PipelineStatus.PENDING, PipelineStatus.PROCESSING]:
            return Response({'error': 'Document not in assignable state'}, status=400)
            
        from apps.processing.services.core import AssignmentService
        try:
            pages_assigned = AssignmentService.assign_pages()
            doc.refresh_from_db()
            return Response({
                'status': 'triggered', 
                'pages_recently_assigned_system_wide': pages_assigned,
                'pipeline_status': doc.pipeline_status
            })
        except Exception as e:
            return Response({'error': str(e)}, status=400)

    @action(detail=True, methods=['get'])
    def pages(self, request, pk=None):
        doc = self.get_object()
        pages = doc.pages.all().order_by('page_number')
        
        # Verify no gaps
        page_numbers = list(pages.values_list('page_number', flat=True))
        expected = list(range(1, (doc.total_pages or 0) + 1))
        missing = sorted(set(expected) - set(page_numbers))
        
        # We use a nested serializer or just the PageSerializer
        from apps.documents.serializers import PageSerializer
        
        return Response({
            'doc_ref': doc.doc_ref,
            'total_pages': doc.total_pages,
            'pages_found': pages.count(),
            'missing_pages': missing,
            'pages': PageSerializer(pages, many=True).data
        })

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAdminUser])
    def approve(self, request, pk=None):
        doc = self.get_object()
        from common.enums import PipelineStatus
        from apps.processing.services.merge import MergeService
        
        if doc.pipeline_status in [PipelineStatus.MERGED, PipelineStatus.APPROVED] and doc.final_file:
            return Response({'message': 'Document already merged/approved and file is ready.'})
            
        if doc.pipeline_status not in [PipelineStatus.ALL_SUBMITTED, PipelineStatus.MERGED, PipelineStatus.APPROVED, PipelineStatus.READY_TO_ASSIGN, PipelineStatus.IN_PROGRESS]:
            return Response({'error': f'Document in state {doc.pipeline_status}. All pages must be submitted/approved first.'}, status=400)
            
        try:
            # Auto-approve all pending submissions for this document if triggering global approve
            from apps.processing.models import SubmittedPage
            from common.enums import ReviewStatus
            
            pending_submissions = SubmittedPage.objects.filter(
                document=doc,
                review_status=ReviewStatus.PENDING_REVIEW
            )
            if pending_submissions.exists():
                count = pending_submissions.count()
                pending_submissions.update(
                    review_status=ReviewStatus.APPROVED,
                    reviewed_at=timezone.now(),
                    reviewed_by=request.user,
                    review_notes="Auto-approved via document-level approval."
                )
                logger.info(f"Auto-approved {count} submissions for document {doc.id}")

            MergeService.merge_approved_pages(doc, request.user.id)
            return Response({'status': 'approved_and_merged'})
        except Exception as e:
            logger.error(f"Approval/Merge failed: {e}", exc_info=True)
            return Response({'error': str(e)}, status=400)

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAdminUser])
    def trigger_merge(self, request, pk=None):
        """
        Explicitly trigger the DOCX merge process.

        New Behaviour (bundles-first):
        1. Auto-approve any bundles in SUBMITTED / UPLOADED / COMPLETED state.
        2. Merge all bundles that have an output_file in strict bundle_index order
           using StrictBundleMergeEngine — regardless of whether a job_id exists.
        3. Fall back to per-page MergeService ONLY when no bundles with output_file
           are found (pure legacy page-based flow).
        """
        document = self.get_object()
        from apps.desktop_bridge.models import AssignmentBundle
        from django.utils import timezone

        try:
            # ── Guard: already merged ─────────────────────────────────
            if document.final_word_file:
                document.refresh_from_db()
                if document.final_word_file:
                    return Response({
                        'status': 'ALREADY_MERGED',
                        'success': True,
                        'message': 'Document is already merged.',
                        'download_url': document.final_word_file.url,
                        'generated_at': document.final_word_generated_at
                    })

            # ── Step 1: Auto-approve all ready bundles ─────────────────
            ready_statuses = [
                AssignmentBundle.Status.SUBMITTED,
                AssignmentBundle.Status.UPLOADED,
                AssignmentBundle.Status.COMPLETED,
                AssignmentBundle.Status.EDITING,
                AssignmentBundle.Status.DOWNLOADED,
            ]
            approved_count = AssignmentBundle.objects.filter(
                document=document,
                status__in=ready_statuses
            ).update(
                status=AssignmentBundle.Status.APPROVED,
                reviewed_at=timezone.now(),
                reviewed_by=request.user
            )
            if approved_count:
                logger.info(f"trigger_merge: Auto-approved {approved_count} bundles for doc {document.id}")

            # ── Step 2: Check if any bundles have uploaded files (new flow) ──
            from django.db.models import Q
            bundles_with_output = AssignmentBundle.objects.filter(
                document=document,
                status__in=[AssignmentBundle.Status.APPROVED, AssignmentBundle.Status.COMPLETED]
            ).filter(
                Q(output_file__isnull=False) & ~Q(output_file='') |
                Q(result_pdf__isnull=False) & ~Q(result_pdf='')
            ).order_by('bundle_index')

            if bundles_with_output.exists():
                # Use StrictBundleMergeEngine — reads output_file directly; merge.py fallback reads result_pdf

                from apps.desktop_bridge.services import StrictBundleMergeEngine
                engine = StrictBundleMergeEngine(document.id)
                success, message = engine.execute()

                if success:
                    document.refresh_from_db()
                    return Response({
                        'status': 'SUCCESS',
                        'success': True,
                        'message': f'Merged {bundles_with_output.count()} bundles successfully.',
                        'download_url': document.final_word_file.url if document.final_word_file else '',
                        'generated_at': document.final_word_generated_at
                    })
                else:
                    return Response({
                        'status': 'VALIDATION_ERROR',
                        'success': False,
                        'error': message
                    }, status=400)

            # ── Step 3: Legacy fallback — per-page SubmittedPage merge ──
            from apps.processing.models import SubmittedPage
            from common.enums import ReviewStatus
            SubmittedPage.objects.filter(
                document=document,
                review_status=ReviewStatus.PENDING_REVIEW
            ).update(
                review_status=ReviewStatus.APPROVED,
                reviewed_at=timezone.now(),
                reviewed_by=request.user,
                review_notes="Auto-approved via administrative merge trigger."
            )

            from apps.processing.services.merge import MergeService
            download_url = MergeService.merge_approved_docx_pages(document, admin_user_id=request.user.id)

            return Response({
                'status': 'SUCCESS',
                'success': True,
                'message': 'Legacy page-by-page merge completed.',
                'download_url': download_url,
                'generated_at': document.final_word_generated_at
            })

        except ValueError as e:
            return Response({
                'status': 'VALIDATION_ERROR',
                'success': False,
                'error': str(e)
            }, status=400)
        except Exception as e:
            logger.error(f"DOCX Merge failed for {document.id}: {e}", exc_info=True)
            return Response({
                'status': 'SERVER_ERROR',
                'success': False,
                'error': f'An internal error occurred during merge: {str(e)}'
            }, status=500)


    @action(detail=True, methods=['get'], url_path='download-final')
    def download_final(self, request, pk=None):
        """
        GET /documents/{id}/download-final/
        
        Strictly deterministic download for the ZeroLossMergeEngine output.
        Only returns the file if pipeline_status is MERGED.
        """
        document = self.get_object()
        
        from common.enums import PipelineStatus
        if document.pipeline_status != PipelineStatus.COMPLETED:
            return Response({
                "error": "Document merge is not complete.",
                "pipeline_status": document.pipeline_status,
                "detail": "The final DOCX can only be downloaded after a successful ZeroLossMergeEngine execution."
            }, status=status.HTTP_400_BAD_REQUEST)
            
        if not document.final_word_file:
            return Response({
                "error": "Final DOCX file missing from storage.",
                "detail": "The merge record indicates completion, but the physical file is not found."
            }, status=status.HTTP_404_NOT_FOUND)
            
        try:
            # Using FileResponse for direct, streaming, efficient download
            from django.http import FileResponse
            from pathlib import Path
            
            # Use original filename from storage if available, otherwise document name
            file_name_raw = Path(document.final_word_file.name).name
            
            response = FileResponse(
                document.final_word_file.open('rb'),
                as_attachment=True,
                filename=file_name_raw,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            return response
            
        except Exception as e:
            return Response({
                "error": "Failed to serve download header.",
                "detail": str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['get'])
    def download(self, request, pk=None, filename=None):
        """
        Download the final merged Word file.
        Prefers the pre-generated 'final_word_file' if available.
        """
        document = self.get_object()

        # 1. Prefer the robustly merged file if it exists
        if document.final_word_file:
            try:
                response = HttpResponse(
                    document.final_word_file.read(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                disp_attr = f'attachment; filename="{document.final_word_file.name.split("/")[-1]}"'
                response['Content-Disposition'] = disp_attr
                return response
            except Exception as e:
                logger.error(f"Error reading final_word_file for {document.id}: {e}")
                # Fallback to dynamic generation if file is missing from storage

        # 2. Dynamic generation fallback (Legacy/Safety)
        try:
            from io import BytesIO
            buffer, output_name = _build_approved_submissions_word_export(document)
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{output_name}"'
            
            # Audit Log
            from apps.audit.models import AuditLog
            from common.enums import AuditEventType
            AuditLog.objects.create(
                action=AuditEventType.DOC_DOWNLOADED,
                document_id=document.id,
                actor=request.user,
                metadata={
                    'format': 'docx',
                    'filename': output_name,
                    'source': 'dynamic_fallback',
                }
            )

            return response
        except ValueError as e:
            return Response({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f"Word merge download failed for {document.id}: {e}", exc_info=True)
            return Response({'error': str(e)}, status=500)

class PageViewSet(viewsets.ModelViewSet):
    queryset = Page.objects.all()
    serializer_class = PageSerializer
    permission_classes = [permissions.IsAuthenticated]

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAdminUser])
    def reassign(self, request, pk=None):
        """Manually reassign a page to a specific user"""
        user_id = request.data.get('user_id')
        if not user_id:
            return Response({'error': 'user_id required'}, status=400)
            
        from apps.accounts.models import ResourceProfile
        from apps.processing.models import PageAssignment
        from common.enums import PageAssignmentStatus
        from django.db import transaction
        
        page = self.get_object()
        target_res = get_object_or_404(ResourceProfile, user_id=user_id)
        
        with transaction.atomic():
            # 1. Capture old assignments to link them
            old_assignments = list(PageAssignment.objects.filter(
                page=page,
                status__in=[PageAssignmentStatus.ASSIGNED, PageAssignmentStatus.IN_PROGRESS]
            ))
            old_asgn = old_assignments[0] if old_assignments else None
            
            # 2. Cancel existing active assignments
            PageAssignment.objects.filter(
                page=page,
                status__in=[PageAssignmentStatus.ASSIGNED, PageAssignmentStatus.IN_PROGRESS]
            ).update(
                status=PageAssignmentStatus.REASSIGNED,
                processing_end_at=timezone.now()
            )
            
            # 3. Create new assignment
            new_asgn = PageAssignment.objects.create(
                page=page,
                resource=target_res,
                document=page.document,
                status=PageAssignmentStatus.ASSIGNED,
                max_processing_time=600,
                is_reassigned=True,
                reassigned_from=old_asgn
            )
            
            # 4. Create Audit Log Entry
            from apps.processing.models import ReassignmentLog
            from common.enums import RejectionReason
            ReassignmentLog.objects.create(
                original_assignment=old_asgn,
                new_assignment=new_asgn,
                reassigned_by=request.user,
                reason=RejectionReason.QUALITY_FAIL, # Generic for manual reassign
                admin_notes=request.data.get('notes', 'Manual administrative reassignment'),
                previous_resource=old_asgn.resource if old_asgn else None,
                new_resource=target_res
            )
            
            # 5. Update Page
            from common.enums import PageStatus
            page.status = PageStatus.ASSIGNED
            page.current_assignee = target_res.user
            page.save()
            
            # 6. Update Resource Load
            target_res.current_load += 1
            target_res.save(update_fields=['current_load'])
            
            # 7. Update old resource load if applicable
            if old_asgn and old_asgn.resource:
                old_asgn.resource.current_load = max(0, old_asgn.resource.current_load - 1)
                old_asgn.resource.save(update_fields=['current_load'])
            
        return Response({'status': 'reassigned', 'resource': target_res.user.username})

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAdminUser])
    def reject(self, request, pk=None):
        from common.enums import PageStatus, AuditEventType
        from apps.audit.models import AuditLog
        
        page = self.get_object()
        page.status = PageStatus.IMPROPERLY_PROCESSED
        page.is_validated = False
        page.save()
        
        AuditLog.objects.create(
            action=AuditEventType.COMPLETED,
            document_id=page.document.id,
            actor=request.user,
            old_status=PageStatus.COMPLETED,
            new_status=PageStatus.IMPROPERLY_PROCESSED,
            metadata={'reason': request.data.get('reason', 'Quality issues'), 'page_id': page.id}
        )
        return Response({'status': 'rejected'})

class BlockUpdateView(views.APIView):
    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request, block_id):
        """
        Granularly save a single block's text.
        Updates Block.current_text and creates a BlockEdit record.
        """
        block = get_object_or_404(Block, id=block_id)

        # Security Check: Must be the active assignee or admin
        is_admin = request.user.role == UserRole.ADMIN or request.user.is_superuser or request.user.is_staff
        if not is_admin:
            from apps.processing.models import PageAssignment, PageAssignmentStatus
            has_active = PageAssignment.objects.filter(
                page=block.page,
                resource__user=request.user,
                status__in=[PageAssignmentStatus.ASSIGNED, PageAssignmentStatus.IN_PROGRESS]
            ).exists()
            if not has_active:
                return Response({'error': 'Permission denied. Page not assigned or already submitted.'}, status=403)

        new_text = request.data.get('text')
        
        if new_text is None:
            return Response({"error": "No text provided"}, status=status.HTTP_400_BAD_REQUEST)
            
        from django.db import transaction
        with transaction.atomic():
            # 1. Update Block
            block.current_text = new_text
            block.save()
            
            # 2. Create Audit Record
            BlockEdit.objects.create(
                block=block,
                edited_by=request.user,
                text=new_text,
                page_num=block.page.page_number
            )
            
        return Response({"status": "saved", "block_id": block_id})

class ConversionRetryView(views.APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, document_id):
        from apps.documents.models import Document
        from common.enums import ConversionStatus, PipelineStatus
        doc = get_object_or_404(Document, pk=document_id)

        if doc.conversion_status not in [ConversionStatus.CONVERSION_FAILED]:
            return Response(
                {'error': True, 'code': 'INVALID_STATE',
                 'message': 'Document is not in a failed state'},
                status=400
            )

        doc.conversion_status = ConversionStatus.PENDING
        doc.conversion_error = ''
        doc.pipeline_status = PipelineStatus.CONVERTING
        doc.save()

        from apps.documents.tasks import convert_word_to_pdf
        task = convert_word_to_pdf.delay(doc.id)
        doc.celery_task_id = task.id
        doc.save(update_fields=['celery_task_id'])

        return Response({
            'success': True,
            'document_id': document_id,
            'websocket_url': f'/ws/conversion/{document_id}/',
            'message': 'Conversion retry started.',
        })

class ConversionStatusView(views.APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, document_id):
        from apps.documents.models import Document
        doc = get_object_or_404(Document, pk=document_id)
        return Response({
            'document_id': doc.id,
            'conversion_status': doc.conversion_status,
            'pipeline_status': doc.pipeline_status,
            'conversion_error': doc.conversion_error,
            'started_at': doc.conversion_started_at,
            'completed_at': doc.conversion_completed_at,
            'doc_ref': doc.doc_ref,
        })
