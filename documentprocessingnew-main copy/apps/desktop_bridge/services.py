from __future__ import annotations

import hashlib
import io
import json
import zipfile
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from django.core.files.base import ContentFile
from django.core.files import File
from django.db import transaction, models
from django.shortcuts import get_object_or_404
from django.utils import timezone
from pypdf import PdfReader, PdfWriter
from django.conf import settings
from django.core.cache import cache
from docx import Document as DocxDocument
from docxcompose.composer import Composer

from .models import (
    AssignmentBundle, UploadedPDF, PageVersion, 
    MergeManifest, MergeAuditLog
)
from apps.documents.models import Page, Document
from apps.processing.models import SubmittedPage, PageAssignment
from common.enums import PageAssignmentStatus, ReviewStatus, PipelineStatus

logger = logging.getLogger(__name__)


def sha256_fileobj(fileobj) -> str:
    """Calculate SHA256 of an open file object without loading it all into memory."""
    h = hashlib.sha256()
    for chunk in iter(lambda: fileobj.read(1024 * 1024), b""):
        h.update(chunk)
    fileobj.seek(0)
    return h.hexdigest()


def build_bundle_manifest(bundle: AssignmentBundle) -> dict:
    """Generate a manifest for the desktop agent."""
    return {
        "metadata": {
            "app": "DocPro",
            "version": "2.1.0",
            "timestamp": timezone.now().isoformat()
        },
        "bundle": {
            "id": str(bundle.id),
            "job_id": str(bundle.job_id) if bundle.job_id else None,
            "document_id": str(bundle.document_id),
            "user_id": str(bundle.user_id) if bundle.user_id else None,
            "bundle_index": bundle.bundle_index,
            "page_start": bundle.page_start,
            "page_end": bundle.page_end,
            "page_numbers": bundle.page_numbers,
            "lease_token": str(bundle.lease_token),
            "source_sha256": bundle.source_sha256,
            "expiry": bundle.lease_expires_at.isoformat() if bundle.lease_expires_at else None
        }
    }


def make_bundle_zip(bundle: AssignmentBundle) -> bytes:
    """Creates a ZIP containing the manifest and the source PDF."""
    manifest = build_bundle_manifest(bundle)
    bio = io.BytesIO()

    with zipfile.ZipFile(bio, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        # 1. Manifest
        zf.writestr("manifest.json", json.dumps(manifest, indent=2))
        
        # 2. Source PDF
        bundle.source_pdf.open("rb")
        try:
            zf.writestr("bundle.pdf", bundle.source_pdf.read())
        finally:
            bundle.source_pdf.close()

    bio.seek(0)
    return bio.read()


@transaction.atomic
def register_upload(bundle: AssignmentBundle, uploaded_file: File) -> UploadedPDF:
    """
    Production-grade upload registration.
    Entry point for the page-centric ledger system.
    """
    document = bundle.document
    
    # 1. Create the UploadedPDF record
    upload = UploadedPDF.objects.create(
        document=document,
        bundle=bundle,
        device=bundle.leased_to,
        file=uploaded_file,
        checksum=sha256_fileobj(uploaded_file.file)
    )

    # 2. Extract and register individual PageVersions
    reader = PdfReader(upload.file.path)
    page_numbers = bundle.page_numbers # This was recorded during slicing
    
    # Validation: Uploaded PDF page count must match bundle page count
    if len(reader.pages) != len(page_numbers):
        logger.error(f"Upload {upload.id} length mismatch: expected {len(page_numbers)}, got {len(reader.pages)}")
        # We still save the upload for audit, but it won't produce valid versions
        return upload

    versions = []
    for idx, p_num in enumerate(page_numbers):
        page_obj = Page.objects.filter(document=document, page_number=p_num).first()
        if not page_obj:
            continue
            
        versions.append(PageVersion(
            document=document,
            page=page_obj,
            bundle=bundle,
            uploaded_pdf=upload,
            page_number=p_num,
            page_index_in_pdf=idx,
            slice_size=len(page_numbers)
        ))
    
    PageVersion.objects.bulk_create(versions)
    
    MergeAuditLog.objects.create(
        document=document,
        event_type="UPLOAD_REGISTERED",
        details={
            "upload_id": str(upload.id),
            "bundle_id": str(bundle.id),
            "pages": page_numbers,
            "version_count": len(versions)
        }
    )
    
    return upload


@transaction.atomic
def register_page_upload(bundle: AssignmentBundle, page_number: int, uploaded_file: File, bundle_version_id=None) -> PageVersion:
    """
    Registers a single page upload. Simple, no version locking.
    """
    from apps.processing.models import SubmittedPage

    document = bundle.document

    # 1. Create UploadedPDF record
    upload = UploadedPDF.objects.create(
        document=document,
        bundle=bundle,
        device=bundle.leased_to,
        file=uploaded_file,
        checksum=sha256_fileobj(uploaded_file.file)
    )

    # 2. Find the page record
    page_obj = get_object_or_404(Page, document=document, page_number=page_number)

    # 3. Create or update PageVersion (upsert: only one version per page per bundle)
    version, _ = PageVersion.objects.update_or_create(
        document=document,
        bundle=bundle,
        page_number=page_number,
        defaults={
            'page': page_obj,
            'uploaded_pdf': upload,
            'page_index_in_pdf': 0,
            'slice_size': 1,
            'is_valid': True,
        }
    )

    # 4. Auto-submit the assignment for review
    from apps.processing.services.core import ProcessingService
    from apps.processing.models import PageAssignment
    from common.enums import PageAssignmentStatus

    assignment = PageAssignment.objects.filter(
        document=document,
        page=page_obj,
        status__in=[PageAssignmentStatus.ASSIGNED, PageAssignmentStatus.IN_PROGRESS]
    ).first()

    if assignment:
        user = bundle.leased_to.user if bundle.leased_to else document.client
        try:
            ProcessingService.complete_assignment(
                assignment.id,
                user,
                uploaded_file=uploaded_file,
                bundle=bundle,
            )
        except Exception as e:
            logger.warning(f"Auto-submit for page {page_number} failed (non-fatal): {e}")

    MergeAuditLog.objects.create(
        document=document,
        event_type="PAGE_UPLOADED",
        details={
            "upload_id": str(upload.id),
            "bundle_id": str(bundle.id),
            "page_number": page_number,
            "filename": uploaded_file.name
        }
    )

    return version


def recompute_bundle_status(bundle_id: uuid.UUID) -> AssignmentBundle.Status:
    """
    Atomic check to see if a bundle is COMPLETED.
    A bundle is COMPLETED iff:
    1. All pages in bundle.page_numbers have an APPROVED Submission.
    2. All submissions belong to the SAME resource (user).
    """
    bundle = AssignmentBundle.objects.get(id=bundle_id)
    page_numbers = set(bundle.page_numbers)
    
    submissions = SubmittedPage.objects.filter(
        bundle=bundle,
        page_number__in=page_numbers
    ).select_related('submitted_by')
    
    if submissions.count() < len(page_numbers):
        return bundle.status # Not all pages submitted yet
        
    # Check for approval and single-user consistency
    first_user = None
    all_approved = True
    page_hits = set()
    
    for sub in submissions:
        if sub.review_status != ReviewStatus.APPROVED:
            all_approved = False
            break
        
        if first_user is None:
            first_user = sub.submitted_by_id
        elif first_user != sub.submitted_by_id:
            # MIXED BUNDLE: This is a violation of the atomic unit rule.
            # We fail the bundle if it contains mixed versions from different users.
            logger.error(f"Bundle {bundle_id} has mixed submissions: User {first_user} and {sub.submitted_by_id}")
            bundle.status = AssignmentBundle.Status.FAILED
            bundle.save()
            return bundle.status

        page_hits.add(sub.page_number)

    if all_approved and page_hits == page_numbers:
        bundle.status = AssignmentBundle.Status.COMPLETED
        bundle.save()
        logger.info(f"Bundle {bundle_id} marked as COMPLETED.")
        
        # Trigger document merge check
        from .tasks import task_deterministic_merge
        task_deterministic_merge.delay(str(bundle.document_id))
        
    return bundle.status


class PageVersionResolver:
    """
    Deterministic resolution logic for selecting the best PageVersion for every slot.
    """
    def __init__(self, document):
        self.document = document

    def resolve(self) -> Dict[int, PageVersion]:
        """
        Returns a map of page_number -> best PageVersion.
        Rules: Smallest slice_size -> Newest updated_at -> UUID tie-breaker.
        """
        # Get all valid versions for this document
        versions = PageVersion.objects.filter(
            document=self.document, 
            is_valid=True
        ).select_related('uploaded_pdf').order_by(
            'page_number',
            'slice_size',      # Smallest first
            '-updated_at',     # Newest first
            'id'               # Final stable tie-breaker
        )

        resolved_map = {}
        for version in versions:
            if version.page_number not in resolved_map:
                resolved_map[version.page_number] = version
        
        return resolved_map


class ZeroLossMergeEngine:
    """
    The orchestrator for document reconstruction.
    Guarantees deterministic assembly and full coverage.
    """
    def __init__(self, document):
        self.document = document
        self.resolver = PageVersionResolver(document)

    @transaction.atomic
    def execute(self) -> MergeManifest:
        manifest = MergeManifest.objects.create(document=self.document)
        
        try:
            # 1. Resolve every page
            page_map = self.resolver.resolve()
            total_pages = self.document.total_pages or 0
            
            # 2. Check for missing pages
            missing = []
            for p_num in range(1, total_pages + 1):
                if p_num not in page_map:
                    missing.append(p_num)
            
            if missing:
                manifest.status = MergeManifest.Status.PARTIAL
                manifest.missing_pages = missing
                manifest.error_details = f"Missing continuity for pages: {missing}"
                manifest.save()
                
                MergeAuditLog.objects.create(
                    document=self.document,
                    manifest=manifest,
                    event_type="MERGE_INCOMPLETE",
                    details={"missing": missing}
                )
                return manifest

            # 3. Deterministic Reconstruction
            writer = PdfWriter()
            version_summary = {}

            # We pre-open readers to avoid excessive I/O
            readers = {} # uploaded_pdf_id -> PdfReader
            
            for p_num in range(1, total_pages + 1):
                version = page_map[p_num]
                pdf_id = str(version.uploaded_pdf.id)
                
                if pdf_id not in readers:
                    readers[pdf_id] = PdfReader(version.uploaded_pdf.file.path)
                
                reader = readers[pdf_id]
                writer.add_page(reader.pages[version.page_index_in_pdf])
                version_summary[p_num] = str(version.id)

            # 4. Save Final PDF
            output_dir = Path(settings.MEDIA_ROOT) / "desktop_bridge" / "merged" / timezone.now().strftime("%Y/%m/%d")
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"reconstructed_{self.document.id}.pdf"

            with open(output_path, "wb") as f:
                writer.write(f)

            # Update Manifest
            manifest.status = MergeManifest.Status.SUCCESS
            manifest.version_map = version_summary
            manifest.completed_at = timezone.now()
            
            # Save the file (relative to media root)
            rel_path = output_path.relative_to(settings.MEDIA_ROOT)
            manifest.final_pdf.name = str(rel_path)
            manifest.save()

            MergeAuditLog.objects.create(
                document=self.document,
                manifest=manifest,
                event_type="MERGE_SUCCESS",
                details={"version_map": version_summary}
            )

        except Exception as e:
            logger.exception(f"Deterministic merge failed for {self.document.id}")
            manifest.status = MergeManifest.Status.FAILED
            manifest.error_details = str(e)
            manifest.save()
            
            MergeAuditLog.objects.create(
                document=self.document,
                manifest=manifest,
                event_type="MERGE_FAILED",
                details={"error": str(e)}
            )
            raise e

        return manifest


class BundleMergeEngine:
    """
    Production-grade Bundle-Based Deterministic Merge System.
    Reconstructs the final DOCX by iterating through COMPLETED bundles.
    """
    def __init__(self, document: Document):
        self.document = document

    @transaction.atomic
    def execute(self) -> Tuple[bool, str]:
        """
        Executes the bundle-based merge logic with strict concurrency safety and idempotency.
        Returns: (success_bool, message)
        """
        # 1. CONCURRENCY LOCK (CRITICAL)
        doc = Document.objects.select_for_update().get(id=self.document.id)
        
        # 2. IDEMPOTENCY CHECK
        if doc.final_word_file and doc.pipeline_status == PipelineStatus.COMPLETED:
            return True, "Document already merged successfully."

        # 3. Fetch COMPLETED bundles
        bundles = list(AssignmentBundle.objects.filter(
            document=doc
        ).order_by('page_start'))

        if not bundles:
            return False, "No bundles found for this document."

        # 4. RESOLVE WINNING SUBMISSIONS (FIRST COMPLETE VALID VERSION PER BUNDLE)
        total_pages = doc.total_pages or 0
        expected_start = 1
        manifest_data = []
        winning_submissions_by_page = {} # page_number -> SubmittedPage
        
        for b in bundles:
            # Rule: Sequential Continuity (PDF Order is Truth)
            if b.page_start != expected_start:
                return False, f"Merge blocked: Gap detected. Expected page {expected_start}, but bundle starts at {b.page_start}."
            
            # Find Winning Version for this Bundle
            # Resolver Rule: 1. completeness, 2. approval state, 3. earliest submitted_at, 4. earliest id
            # We group by bundle_version_id to ensure atomicity
            candidate_versions = SubmittedPage.objects.filter(
                bundle=b
            ).values('bundle_version_id').annotate(
                page_count=models.Count('page_number', distinct=True)
            ).filter(page_count=len(b.page_numbers))
            
            winning_version_id = None
            earliest_v_time = None
            
            for candidate in candidate_versions:
                vid = candidate['bundle_version_id']
                v_pages = SubmittedPage.objects.filter(bundle_version_id=vid, bundle=b)
                
                # Check if ALL pages in this version are APPROVED
                if v_pages.filter(review_status=ReviewStatus.APPROVED).count() == len(b.page_numbers):
                    # Pick the submitted_at from the latest page in the version for sorting
                    # Or better: pick from the related assignment
                    first_sub = v_pages.first()
                    v_time = first_sub.assignment.submitted_at or first_sub.assignment.assigned_at
                    
                    if winning_version_id is None or v_time < earliest_v_time or (v_time == earliest_v_time and str(vid) < str(winning_version_id)):
                        winning_version_id = vid
                        earliest_v_time = v_time
            
            if not winning_version_id:
                return False, f"Merge blocked: Bundle {b.bundle_index} ({b.page_start}-{b.page_end}) has no winning (fully approved) submission."

            # Lock the winner for this bundle
            winner_pages = SubmittedPage.objects.filter(bundle_version_id=winning_version_id, bundle=b).select_related('submitted_by', 'assignment').order_by('page_number')
            for sub in winner_pages:
                # FILE INTEGRITY CHECK
                if not sub.output_page_file or sub.output_page_file.size == 0:
                     return False, f"Merge blocked: Winning version for Page {sub.page_number} has an invalid or empty file."
                
                winning_submissions_by_page[sub.page_number] = sub
                
                # Gather manifest info (showing exactly which version was chosen)
                manifest_data.append({
                    "page_number": sub.page_number,
                    "bundle_id": str(b.id),
                    "bundle_version_id": str(winning_version_id),
                    "user_id": str(sub.submitted_by_id),
                    "username": sub.submitted_by.username,
                    "submitted_at": earliest_v_time.isoformat() if earliest_v_time else None,
                    "file_hash": self._get_file_hash(sub.output_page_file),
                    "source_pdf_page_range": f"{b.page_start}-{b.page_end}"
                })

            expected_start = b.page_end + 1

        # Final Continuity Check against original PDF total pages
        if expected_start - 1 != total_pages:
             return False, f"Merge blocked: Document has {total_pages} pages, but bundles only cover up to page {expected_start - 1}."
        
        if len(winning_submissions_by_page) != total_pages:
             return False, f"Merge blocked: Final page set size ({len(winning_submissions_by_page)}) mismatch with document total ({total_pages})."

        # 3. High-Fidelity DOCX Composition
        try:
            merged_docx = self._compose_docx(manifest_data, winning_submissions_by_page)
            
            # 4. Storage & State
            filename = f"DocPro_Final_{doc.id.hex[:8]}.docx"
            temp_path = Path(settings.MEDIA_ROOT) / "temp" / filename
            temp_path.parent.mkdir(parents=True, exist_ok=True)
            
            merged_docx.save(str(temp_path))
            
            from django.core.files import File
            with open(temp_path, "rb") as f:
                doc.final_word_file.save(filename, File(f), save=False)
            
            doc.final_word_manifest = {"pages": manifest_data, "merged_at": timezone.now().isoformat()}
            doc.final_word_generated_at = timezone.now()
            doc.pipeline_status = PipelineStatus.COMPLETED
            doc.save()
            
            # Cleanup
            if temp_path.exists():
                temp_path.unlink()
                
            return True, "Deterministic first-submission bundle merge completed successfully."
        except Exception as e:
            logger.exception(f"DOCX Composition failed for doc {doc.id}: {e}")
            doc.final_word_error = str(e)
            doc.pipeline_status = PipelineStatus.FAILED
            doc.save()
            return False, f"Composition failed: {e}"

    def _get_file_hash(self, file_field) -> str:
        if not file_field: return ""
        try:
            file_field.open('rb')
            h = hashlib.sha256()
            for chunk in iter(lambda: file_field.read(4096), b""):
                h.update(chunk)
            file_field.close()
            return h.hexdigest()
        except Exception:
            return ""

    def _compose_docx(self, manifest_data: List[dict], sub_map: Dict[int, SubmittedPage]) -> DocxDocument:
        """Uses docxcompose to merge the individual page docx files."""
        # Start with the first page
        import io
        first_page_num = manifest_data[0]["page_number"]
        first_sub = sub_map[first_page_num]
        
        first_sub.output_page_file.open('rb')
        master = DocxDocument(io.BytesIO(first_sub.output_page_file.read()))
        first_sub.output_page_file.close()
        
        composer = Composer(master)
        
        # Append subsequent pages
        for i in range(1, len(manifest_data)):
            p_num = manifest_data[i]["page_number"]
            sub = sub_map[p_num]
            
            sub.output_page_file.open('rb')
            next_doc_io = io.BytesIO(sub.output_page_file.read())
            sub.output_page_file.close()
            
            try:
                next_doc_io.seek(0)
                next_doc = DocxDocument(next_doc_io)
                composer.append(next_doc)
            except Exception as compose_error:
                # 🔴 3. DOCX FAILURE FALLBACK
                logger.error(f"docxcompose failed for page {p_num} in document {self.document.id}: {compose_error}")
                try:
                    # Fallback: Simple python-docx append attempt
                    next_doc_io.seek(0)
                    next_doc = DocxDocument(next_doc_io)
                    for element in next_doc.element.body:
                        master.element.body.append(element)
                except Exception as fallback_error:
                    logger.error(f"Fallback merge also failed for page {p_num}: {fallback_error}")
                    raise Exception(f"Fatal merge error on Page {p_num}: {compose_error} (Fallback: {fallback_error})")
            
        return master


def regenerate_bundle_file(bundle: AssignmentBundle) -> bool:
    """
    On-demand regeneration of a bundle's source PDF from the master document.
    Used for self-healing when files are missing from disk.
    """
    try:
        from pypdf import PdfReader, PdfWriter
        import io
        from django.core.files import File

        doc = bundle.document
        source_path = None

        # Priority: internal processing file > original file
        # Check storage existence to avoid FileNotFoundError on .path access
        if doc.file and doc.file.storage.exists(doc.file.name):
            source_path = doc.file.path
        elif doc.original_file and doc.original_file.storage.exists(doc.original_file.name):
            source_path = doc.original_file.path

        if not source_path:
            logger.error(f"Cannot regenerate bundle {bundle.id}: Master document files are missing from storage.")
            return False

        reader = PdfReader(source_path)
        writer = PdfWriter()

        for p_num in bundle.page_numbers:
            # pypdf is 0-indexed, page_numbers are 1-indexed
            writer.add_page(reader.pages[p_num - 1])

        tmp = io.BytesIO()
        writer.write(tmp)
        tmp.seek(0)

        filename = f"bundle_{doc.id}_idx{bundle.bundle_index}.pdf"
        bundle.source_pdf.save(filename, File(tmp), save=True)
        bundle.source_sha256 = sha256_fileobj(tmp)
        bundle.save()

        logger.info(f"Successfully regenerated missing source_pdf for bundle {bundle.id}")
        return True
    except Exception as e:
        logger.exception(f"Failed to regenerate bundle {bundle.id}: {e}")
        return False


def create_bundles_for_document(document, pages_per_bundle=10):
    """
    Logic for partitioning a document into contiguous bundles for the agent.
    Acts as the entry point for Strict Bundle Flow.
    """
    from apps.documents.models import Page
    from apps.processing.models import Job
    from django.core.files import File
    from django.utils import timezone
    
    # NEW: Cleanup old bundles to ensure 'Strict Flow' is correctly applied on retry
    AssignmentBundle.objects.filter(document=document).delete()
    logger.info(f"Cleared old bundles for document {document.id} to enforce fresh Strict Flow.")
    
    # ENFORCE STRICT FLOW: Ensure document has a Job
    if not document.job:
        job = Job.objects.create(
            name=f"Direct Processing: {document.name or document.title}",
            created_at=timezone.now()
        )
        document.job = job
        document.save(update_fields=['job'])
        logger.info(f"Created default Job {job.id} for document {document.id} to enforce Strict Flow.")

    pages = list(Page.objects.filter(document=document).order_by('page_number'))
    if not pages:
        return
    
    total_pages = len(pages)
    bundle_count = (total_pages + pages_per_bundle - 1) // pages_per_bundle
    
    for i in range(bundle_count):
        start_idx = i * pages_per_bundle
        end_idx = min(start_idx + pages_per_bundle, total_pages)
        bundle_pages = pages[start_idx:end_idx]
        
        page_nums = [p.page_number for p in bundle_pages]
        p_start = page_nums[0]
        p_end = page_nums[-1]
        
        # Create Bundle Record
        bundle = AssignmentBundle.objects.create(
            document=document,
            job=document.job,
            bundle_index=i,
            page_start=p_start,
            page_end=p_end,
            page_numbers=page_nums,
            status=AssignmentBundle.Status.READY
        )
        
        # Use healing logic to generate the initial file
        regenerate_bundle_file(bundle)
        
    logger.info(f"Generated {bundle_count} bundles for document {document.id} in STRICT FLOW mode.")
def trigger_merge_if_ready(document_id):
    """
    Checks if all bundles for a document are APPROVED.
    If so, enqueues the strict bundle merge task.
    """
    total_bundles = AssignmentBundle.objects.filter(document_id=document_id).count()
    approved_count = AssignmentBundle.objects.filter(
        document_id=document_id, 
        status=AssignmentBundle.Status.APPROVED
    ).count()

    if total_bundles > 0 and total_bundles == approved_count:
        logger.info(f"All {total_bundles} bundles for Doc {document_id} are APPROVED. Triggering merge.")
        from .tasks import task_deterministic_merge
        task_deterministic_merge.delay(str(document_id))
    else:
        logger.info(f"Doc {document_id} approval: {approved_count}/{total_bundles} bundles APPROVED.")


class StrictBundleMergeEngine:
    """
    Production-grade, deterministic merge engine for Job-based documents.
    Operates ONLY on atomic bundles (DOCX) in bundle_index order.
    """
    def __init__(self, document_id):
        self.document_id = document_id
        self.lock_key = f"strict_merge_lock_{document_id}"

    def execute(self):
        # 1. Acquire Lock
        if not cache.add(self.lock_key, "locked", timeout=600):
            logger.warning(f"Merge already in progress for doc {self.document_id}")
            return False, "Merge in progress"

        try:
            with transaction.atomic():
                doc = Document.objects.select_for_update().get(id=self.document_id)
                from django.db.models import Q
                # Fetch all APPROVED/COMPLETED bundles that have any uploaded file,
                # ordered by page_start for correct sequential page order
                bundles = AssignmentBundle.objects.filter(
                    document=doc,
                    status__in=[AssignmentBundle.Status.APPROVED, AssignmentBundle.Status.COMPLETED]
                ).filter(
                    Q(output_file__isnull=False) & ~Q(output_file='') |
                    Q(result_pdf__isnull=False) & ~Q(result_pdf='')
                ).order_by('page_start')  # CRITICAL: sort by page_start, not bundle_index

                total = bundles.count()
                if total == 0:
                    return False, "No approved bundles with uploaded DOCX files found. Please upload the bundle files via the desktop agent first."

                logger.info(f"StrictBundleMergeEngine: Merging {total} bundles for doc {self.document_id}")

                # 2. Sequential Merge in page order
                import io
                master_doc = None
                composer = None
                
                total_pages = doc.total_pages or 0
                expected_page = 1
                merged_bundles_count = 0

                for b in bundles:
                    # ── CONTINUITY CHECK ──
                    if b.page_start != expected_page:
                        return False, f"Merge gap detected: Expected page {expected_page}, but bundle {b.bundle_index} starts at {b.page_start}. Please ensure all preceding bundles are approved."

                    # Try output_file first, then result_pdf
                    doc_stream = None
                    file_name = ""
                    for file_field, label in [(b.output_file, 'output_file'), (b.result_pdf, 'result_pdf')]:
                        if not file_field:
                            continue
                        try:
                            file_field.open('rb')
                            data = file_field.read()
                            file_name = file_field.name.lower()
                            file_field.close()
                            if data and len(data) > 64:
                                doc_stream = io.BytesIO(data)
                                logger.info(f"  Bundle {b.bundle_index} ({b.page_start}-{b.page_end}): using {label} ({file_name})")
                                break
                        except Exception as read_err:
                            logger.warning(f"StrictBundleMergeEngine: Cannot read {label} for bundle {b.id}: {read_err}")

                    if doc_stream is None:
                        return False, f"Cannot read any file for bundle {b.id} (index {b.bundle_index}, pages {b.page_start}-{b.page_end}). Please ensure the file was uploaded correctly."

                    # ── TYPE CHECK ──
                    if not file_name.endswith('.docx'):
                        return False, f"StrictBundleMergeEngine requires DOCX files for merging. Bundle {b.bundle_index} provided a {file_name.split('.')[-1]} file. Please re-upload as DOCX."

                    if master_doc is None:
                        try:
                            master_doc = DocxDocument(doc_stream)
                            composer = Composer(master_doc)
                        except Exception as e:
                            return False, f"Failed to open Bundle {b.bundle_index} as DOCX: {str(e)}"
                    else:
                        try:
                            doc_to_append = DocxDocument(doc_stream)
                            composer.append(doc_to_append)
                        except Exception as e:
                            return False, f"Failed to append Bundle {b.bundle_index} to DOCX: {str(e)}"

                    
                    merged_bundles_count += 1
                    expected_page = b.page_end + 1

                # ── FINAL CONTINUITY CHECK ──
                if expected_page - 1 != total_pages:
                    return False, f"Merge incomplete: Covered up to page {expected_page-1}, but document has {total_pages} total pages. All bundles must be approved for a full merge."


                # 3. Save Final Output
                final_filename = f"{doc.name or doc.id}_final.docx"
                output_dir = Path(settings.MEDIA_ROOT) / "storage" / "3_final_merged_docx" / timezone.now().strftime("%Y/%m/%d")
                output_dir.mkdir(parents=True, exist_ok=True)

                final_path = output_dir / final_filename
                master_doc.save(str(final_path))

                # Update Document record
                with open(final_path, 'rb') as f:
                    doc.final_word_file.save(final_filename, File(f), save=False)

                doc.pipeline_status = PipelineStatus.COMPLETED
                doc.status = "COMPLETED"
                doc.completed_at = timezone.now()
                doc.save()

                # Mark all bundles as MERGED
                bundles.update(status=AssignmentBundle.Status.MERGED)

                logger.info(f"Strict Merge SUCCESS for Document {self.document_id}. Final file saved.")
                return True, "Success"

        except Exception as e:
            logger.exception(f"Strict Merge FAILED for Document {self.document_id}: {e}")
            return False, str(e)
        finally:
            cache.delete(self.lock_key)

def get_or_create_assignment_bundle(document, page_numbers) -> AssignmentBundle:
    """
    Finds or creates an AssignmentBundle that exactly matches the provided page numbers.
    Used for ensuring resources only download their assigned pages.
    """
    page_numbers = sorted(list(set(page_numbers)))
    if not page_numbers:
        return None

    # Deterministic index to avoid duplicates for the same range
    # We use a high base (100,000) to avoid collision with standard 10-page splits
    range_hash = hashlib.md5(",".join(map(str, page_numbers)).encode()).hexdigest()[:7]
    bundle_index = 100000 + int(range_hash, 16)

    bundle = AssignmentBundle.objects.filter(document=document, bundle_index=bundle_index).first()
    if bundle:
        return bundle

    # Create a new sliced PDF
    writer = PdfWriter()
    source_pdf_path = document.file.path if document.file else document.original_file.path
    reader = PdfReader(source_pdf_path)

    for p_num in page_numbers:
        # page_number is 1-indexed, PdfReader uses 0-indexed
        try:
            writer.add_page(reader.pages[p_num - 1])
        except IndexError:
            logger.warning(f"Page {p_num} not found in document {document.id}")

    # Buffer it
    tmp = io.BytesIO()
    writer.write(tmp)
    tmp.seek(0)

    # Create Bundle Record
    bundle = AssignmentBundle.objects.create(
        document=document,
        bundle_index=bundle_index,
        page_start=page_numbers[0],
        page_end=page_numbers[-1],
        page_numbers=page_numbers,
        status=AssignmentBundle.Status.READY
    )

    filename = f"assigned_bundle_{document.id}_range_{page_numbers[0]}-{page_numbers[-1]}.pdf"
    from django.core.files import File
    bundle.source_pdf.save(filename, File(tmp), save=False)
    bundle.source_sha256 = sha256_fileobj(tmp)
    bundle.save()

    logger.info(f"Dynamically created bundle {bundle.id} for document {document.id}, pages {page_numbers}")
    return bundle


def reassign_bundle(bundle_id: uuid.UUID, manager_user) -> bool:
    """
    Reassigns a REJECTED bundle by resetting it to READY.
    Follows Section 6: ensures the prev user is excluded from picking it up immediately.
    """
    with transaction.atomic():
        bundle = AssignmentBundle.objects.select_for_update().get(id=bundle_id)
        
        # 1. Clear current lease and user
        # We keep attempt_count and rejection_reason as audit trail
        bundle.status = AssignmentBundle.Status.READY
        bundle.leased_to = None
        bundle.user = None # Mandatory: assign to someone else next time
        bundle.lease_expires_at = None
        bundle.save()
        
        # 2. Trigger global PageAssignment reassignment for pages in this bundle
        # This keeps the Page Ledger system in sync.
        from apps.processing.models import PageAssignment
        from apps.processing.services.core import AssignmentService
        
        # For simplicity, we mark the latest SUBMITTED assignments for these pages as REJECTED
        assignments = PageAssignment.objects.filter(
            document=bundle.document,
            page__page_number__in=bundle.page_numbers,
            status=PageAssignmentStatus.SUBMITTED
        )
        
        for a in assignments:
            # This resets the page and makes it available for the standard assign_pages task
            AssignmentService.reassign_rejected_assignment(a.id, manager_user, auto_assign=False)
            
        logger.warning(f"Bundle {bundle_id} successfully reset to READY and reassigned.")
        
    return True
